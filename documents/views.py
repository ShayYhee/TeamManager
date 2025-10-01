from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.models import User
from django.contrib.admin.models import LogEntry, CHANGE, ADDITION, DELETION
from django.contrib.auth import login, get_user_model, logout
from django.contrib.auth.signals import user_logged_in
from django.contrib.auth.tokens import default_token_generator
from django.contrib.auth.views import LoginView
from django.contrib.contenttypes.models import ContentType
from django.contrib import messages
from django.core.mail import send_mail, EmailMessage, get_connection
from django.core.paginator import Paginator
from django.core.exceptions import ValidationError, PermissionDenied, BadRequest
from django.conf import settings
from django.db import transaction
from django.db.models import Q, F
from django.dispatch import receiver
from django.forms import formset_factory
from django.http import HttpResponse, HttpResponseForbidden, JsonResponse, HttpResponseRedirect, HttpResponseNotFound
from django.urls import reverse
from django.utils.encoding import force_bytes, force_str
from django.utils.http import urlsafe_base64_decode
from django.utils.text import slugify
from django.utils import timezone
from django.utils.timezone import now
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from .forms import DocumentForm, SignUpForm, CreateDocumentForm, FileUploadForm, FolderForm, TaskForm, ReassignTaskForm, StaffProfileForm, StaffDocumentForm, EmailConfigForm, UserForm, DepartmentForm, TeamForm, EventForm, EventParticipantForm, NotificationForm, UserNotificationForm, CompanyProfileForm, ContactForm, EmailForm, EditUserForm, CompanyDocumentForm, SupportForm, VacancyForm, VacancyApplicationForm
from .models import Document, CustomUser, Role, File, Folder, Task, StaffProfile, Notification, UserNotification, StaffDocument, Event, EventParticipant, Department, Team, CompanyProfile, Contact, Email, Attachment, CompanyDocument, Vacancy, VacancyApplication
from raadaa.settings import ALLOWED_HOSTS
from .serializers import EventSerializer
from .placeholders import replace_placeholders
from docx import Document as DocxDocument
from docx.shared import Inches
from ckeditor_uploader.views import upload as ckeditor_upload
import csv, subprocess, platform, shutil, pdfkit, re, smtplib, os, requests, io, urllib.parse, json, logging, imaplib, email
from raadaa import settings
from html2docx import html2docx
from docx2txt import process
from datetime import datetime, date, timedelta
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from django.http import HttpResponse
from rest_framework import viewsets, permissions
from .forms import ForgotPasswordForm, ResetPasswordForm

logger = logging.getLogger(__name__)

User = get_user_model()

main_superuser = CustomUser.objects.filter(is_superuser=True).first()

SUPERUSER_EMAIL_PROVIDER = main_superuser.email_provider
SUPERUSER_EMAIL_ADDRESS = main_superuser.email_address
SUPERUSER_EMAIL_PASSWORD = main_superuser.email_password


def is_admin(user):
    # Check if the user is an admin

    for role in user.roles.all():
        if role.name == "Admin":
            return True

def is_hr(user):
    # Check if the user is an admin

    for role in user.roles.all():
        if role.name == "HR":
            return True

def custom_404(request, exception):
    return render(request, '404.html', {'subdomain': request.get_host().split('.')[0], 'message': str(exception)}, status=404)

def custom_403(request, exception):
    return render(request, '403.html', {'message': str(exception)}, status=403)

def custom_400(request, exception):
    return render(request, '400.html', {'message': str(exception)}, status=400)

def custom_500(request):
    return render(request, '500.html', status=500)

# def mail_connection(sender_email, sender_password):
#     return get_connection(
#         backend="django.core.mail.backends.smtp.EmailBackend",
#         host="smtp.zoho.com",
#         port=587,
#         username=sender_email,
#         password=sender_password,
#         use_tls=True,
#     )

def get_tenant_url(request):
    if hasattr(request, 'user') and request.user.is_authenticated:
        if not CustomUser.objects.filter(id=request.user.id, tenant=request.tenant).exists():
            print(f"User {request.user.username} not associated with tenant {request.tenant.slug if request.tenant else 'None'}")
            expected_subdomain = (
                request.user.tenant.slug
                if hasattr(request.user, 'tenant') and request.user.tenant
                else None
            )
            if expected_subdomain is None:
                logout(request)
                raise PermissionDenied("You have no associated tenant. Contact support. faith.osebi@transnetcloud.com")
            print(f"Wrong user tenant slug: {expected_subdomain}")
            base_domain = "localhost:8000" if settings.DEBUG else "teammanager.ng"
            protocol = "http" if settings.DEBUG else "https"
            home_url = f"{protocol}://{expected_subdomain}.{base_domain}/"
            print(f"Redirecting to tenant home: {home_url}")
            return home_url

def get_email_smtp_connection(sender_provider, sender_email, sender_password):
    smtp_settings = {
        "gmail": ("smtp.gmail.com", 587, True, False),
        "zoho": ("smtp.zoho.com", 587, True, False),
        "yahoo": ("smtp.mail.yahoo.com", 587, True, False),
        "outlook": ("smtp-mail.outlook.com", 587, True, False),
        "icloud": ("smtp.mail.me.com", 587, True, False),
    }
    
    sender_provider = sender_provider.lower()
    if sender_provider not in smtp_settings:
        print(f"Unsupported email provider: {sender_provider}")
        return None, f"Unsupported email provider: {sender_provider}"
    
    host, port, use_tls, use_ssl = smtp_settings[sender_provider]
    try:
        print(f"Connecting to {host}:{port} for {sender_email}")
        connection = get_connection(
            backend="django.core.mail.backends.smtp.EmailBackend",
            host=host,
            port=port,
            username=sender_email,
            password=sender_password,
            use_tls=use_tls,
        )
        # Test the connection by opening it
        connection.open()
        # connection.close()
        print(f"SMTP connection successful for {sender_provider}")
        return connection, None  # Success: return connection and no error
    except smtplib.SMTPAuthenticationError as e:
        print(f"Authentication failed for {sender_provider}: {str(e)}")
        return None, str(e)  # Return error message for user feedback
    except Exception as e:
        print(f"SMTP connection failed for {sender_provider}: {str(e)}")
        return None, str(e)

    
def get_email_imap_connection(sender_provider, sender_email, sender_password):
    if sender_provider == "gmail":
        return imaplib.IMAP4_SSL("imap.gmail.com", 993)
    elif sender_provider == "yahoo":
        return imaplib.IMAP4_SSL("imap.mail.yahoo.com", 993)
    elif sender_provider == "outlook":
        return imaplib.IMAP4_SSL("outlook.office365.com", 993)
    elif sender_provider == "zoho":
        return imaplib.IMAP4_SSL("imap.zoho.com", 993)
    elif sender_provider == "icloud":
        return imaplib.IMAP4_SSL("imap.mail.me.com", 993)
    else:
        return None

class CustomLoginView(LoginView):
    def form_valid(self, form):
        # Authenticate and log in the user
        response = super().form_valid(form)  # Logs in the user
        user = self.request.user

        # If superuser, allow staying on base domain
        if user.is_superuser:
            return response

        # For non-superusers, redirect to their tenant's login URL
        expected_subdomain = (
            user.tenant.slug
            if hasattr(user, 'tenant') and user.tenant
            else None
        )
        if expected_subdomain is None:
            return HttpResponseForbidden("You are not associated with this company. Please ensure your subdomain is correct or contact faith.osebi@transnetcloud.com")
        base_domain = "localhost:8000" if settings.DEBUG else "teammanager.ng"
        protocol = "http" if settings.DEBUG else "https"
        tenant_login_url = f"{protocol}://{expected_subdomain}.{base_domain}/accounts/login"

        # Redirect to tenant-specific login URL
        return redirect(tenant_login_url)

    def get(self, request, *args, **kwargs):
        # If already authenticated, redirect to home or tenant URL
        if request.user.is_authenticated:
            if request.user.is_superuser:
                return redirect(settings.LOGIN_REDIRECT_URL or '/')
            expected_subdomain = (
                request.user.tenant.slug
                if hasattr(request.user, 'tenant') and request.user.tenant
                else None
            )
            if expected_subdomain is None:
                return HttpResponseForbidden("You are not associated with this company. Please ensure your subdomain is correct or contact faith.osebi@transnetcloud.com")
            base_domain = "localhost:8000" if settings.DEBUG else "teammanager.ng"
            protocol = "http" if settings.DEBUG else "https"
            return redirect(f"{protocol}://{expected_subdomain}.{base_domain}/")
        return super().get(request, *args, **kwargs)
    
def forgot_password(request):
    if request.method == 'POST':
        form = ForgotPasswordForm(request.POST)
        if form.is_valid():
            form.save(request)
            return redirect('password_reset_sent')  # Or a 'email sent' page if you want to add one
    else:
        form = ForgotPasswordForm()
    return render(request, 'registration/forgot_password.html', {'form': form})

def reset_password(request, uidb64, token):
    try:
        uid = force_str(urlsafe_base64_decode(uidb64))
        user = User.objects.get(pk=uid)
    except (TypeError, ValueError, OverflowError, User.DoesNotExist):
        user = None

    if user and default_token_generator.check_token(user, token):
        if request.method == 'POST':
            form = ResetPasswordForm(user, request.POST)
            if form.is_valid():
                form.save()
                return redirect('password_reset_success')
        else:
            form = ResetPasswordForm(user)
        return render(request, 'registration/reset_password.html', {'form': form})
    else:
        # Invalid link (expired or tampered)
        return render(request, 'registration/reset_password.html', {'error': 'Invalid or expired reset link.'})

def password_reset_success(request):
    return render(request, 'registration/password_reset_success.html')

def password_reset_sent(request):
    return render(request, 'registration/password_reset_sent.html')

def upload_to_documents_word(document):
    tenant_name = document.tenant.name
    username = document.created_by.username if document.created_by else "anonymous"
    return os.path.join('documents', tenant_name, username, 'word')

def upload_to_documents_pdf(document):
    tenant_name = document.tenant.name
    username = document.created_by.username if document.created_by else "anonymous"
    return os.path.join('documents', tenant_name, username, 'pdf')

def send_approval_request(document):
    # Get all BDM emails for the document's tenant
    bdm_emails = CustomUser.objects.filter(
        tenant=document.tenant,  # Filter by the document's tenant
        roles__name="BDM"
    ).values_list("email", flat=True)

    # Ensure there are BDMs to notify
    if not bdm_emails:
        return
    
    # Ensure the document's creator is from the same tenant
    if document.created_by.tenant != document.tenant:
        # return HttpResponseForbidden("Unauthorized: Creator does not belong to the document's tenant.")
        return render('error.html', {
        'error_code': '403',
        'message': 'Unauthorized: User does not belong to the document\'s company.'
    }, status=403)

    sender_provider = document.created_by.email_provider
    sender_email = document.created_by.email_address
    sender_password = document.created_by.email_password

    if not sender_email or not sender_password:
        return HttpResponseForbidden("Your email credentials are missing. Contact admin.")

    connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)

    subject = f"Approval Request: {document.company_name}"
    message = f"""
    Dear BDM,

    A new document for {document.company_name} has been created by {document.created_by.get_full_name()}.
    Please review and approve it.

    Best regards,  
    {document.created_by.get_full_name()}
    """

    print("Sending mail...")

    # Create email with attachment
    email = EmailMessage(subject, message, sender_email, list(bdm_emails), connection=connection)
    email.attach_file(document.pdf_file.path)  # Attach the PDF
    email.send()
    
    print("Mail Sent")


@login_required
def create_document(request):
    if not hasattr(request, 'tenant') or not request.user.tenant == request.tenant:
        # return HttpResponseForbidden("You are not authorized to perform actions for this tenant.")
        tenant_url = get_tenant_url(request)
        return render(request, 'tenant_error.html', {'error_code': '401', 'message': 'Access denied.', 'user': request.user, 'tenant_url': tenant_url,}) 
    
    if request.user.tenant.slug not in ["raadaa", "transnet-cloud"]:
        # return HttpResponseForbidden("Unauthorized: User can not view this page.")
        raise PermissionDenied()
    
    DocumentFormSet = formset_factory(DocumentForm, extra=1)
    
    if request.method == "POST":
        print("POST request received")
        formset = DocumentFormSet(request.POST, request.FILES)
        print("Formset data:", request.POST)
        print("Files:", request.FILES)

        if formset.is_valid():
            print("Formset is valid")
            for form in formset:
                if form.has_changed():
                    print("Processing form:", form.cleaned_data)
                    document = form.save(commit=False)
                    document.created_by = request.user
                    document.tenant = request.tenant  # Set tenant from middleware
                    # Validate that the user belongs to the same tenant
                    if document.created_by.tenant != request.tenant:
                        return HttpResponse("Unauthorized: User does not belong to the current company.", status=403)

                    # Get or create the "Template Document" folder
                    user = request.user
                    department = user.department if hasattr(user, 'department') else None
                    team = user.teams.first() if hasattr(user, 'teams') and user.teams.exists() else None

                    # Ensure user has a department or team
                    if not department and not team:
                        return HttpResponse("Error: User must be associated with a department or team.", status=403)

                    # Choose department or team for the folder (prefer department if available)
                    folder_defaults = {
                        'created_by': user,
                        'is_public': True,
                    }

                    try:
                        template_folder, created = Folder.objects.get_or_create(
                            tenant=request.tenant,
                            name="Template Document",
                            defaults=folder_defaults
                        )
                    except ValidationError as e:
                        print(f"Folder creation error: {e}")
                        return HttpResponse(f"Error creating Template Document folder: {e}", status=400)
                        # raise BadRequest(f"Error creating Template Document folder: {e}")

                    # Validate folder access (similar to create_public_folder)
                    # if template_folder.department and template_folder.department != user.department:
                    #     return HttpResponse("Invalid Department for Template Document folder.", status=403)
                    # if template_folder.team and template_folder.team not in user.teams.all():
                    #     return HttpResponse("Invalid Team for Template Document folder.", status=403)

                    document.save()

                    creation_method = form.cleaned_data['creation_method']
                    word_dir = os.path.join(settings.MEDIA_ROOT, upload_to_documents_word(document))
                    pdf_dir = os.path.join(settings.MEDIA_ROOT, upload_to_documents_pdf(document))
                    os.makedirs(word_dir, exist_ok=True)
                    os.makedirs(pdf_dir, exist_ok=True)

                    base_filename = f"{document.company_name}_{document.id}"

                    if creation_method == 'template':
                        template_filename = "Approval Letter Template.docx" if document.document_type == "approval" else "SLA Template.docx"
                        template_path = os.path.join(settings.BASE_DIR, "documents/templates/docx_templates", template_filename)
                        if not os.path.exists(template_path):
                            print(f"Template not found: {template_path}")
                            return HttpResponse(f"Error: Template not found at {template_path}", status=500)
                        
                        today = datetime.today()
                        if document.document_type == "approval":
                            day = today.day
                            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
                            formatted_date = today.strftime("%d") + suffix + " " + today.strftime("%d %B, %Y")
                        else:
                            formatted_date = today.strftime("%m/%d/%Y")

                        doc = DocxDocument(template_path)
                        replacements = {
                            "{{Company Name}}": document.company_name,
                            "{{Company Address}}": document.company_address,
                            "{{Contact Person Name}}": document.contact_person_name,
                            "{{Contact Person Email}}": document.contact_person_email,
                            "{{Contact Person Designation}}": document.contact_person_designation + ",",
                            "{{Sales Rep}}": document.sales_rep,
                            "{{Date}}": formatted_date,
                        }
                        doc = replace_placeholders(doc, replacements, document.document_type)

                        word_filename = f"{base_filename}.docx"
                        word_path = os.path.join(word_dir, word_filename)
                        doc.save(word_path)
                        document.word_file = os.path.join(upload_to_documents_word(document), word_filename)

                        # Save Word file to File
                        public_word_file = File(
                            tenant=request.tenant,
                            original_name=word_filename,
                            file=os.path.join(upload_to_documents_word(document), word_filename),
                            folder=template_folder,
                            uploaded_by=user
                        )
                        public_word_file.save()
                    else:
                        uploaded_file = form.cleaned_data['uploaded_file']
                        file_extension = uploaded_file.name.lower().split('.')[-1]

                        if file_extension == 'docx':
                            word_filename = f"{base_filename}.docx"
                            word_path = os.path.join(word_dir, word_filename)
                            with open(word_path, 'wb') as f:
                                for chunk in uploaded_file.chunks():
                                    f.write(chunk)
                            document.word_file = os.path.join(upload_to_documents_word(document), word_filename)

                            # Save Word file to File
                            public_word_file = File(
                                tenant=request.tenant,
                                original_name=word_filename,
                                file=os.path.join(upload_to_documents_word(document), word_filename),
                                folder=template_folder,
                                uploaded_by=user
                            )
                            public_word_file.save()
                        elif file_extension == 'pdf':
                            pdf_filename = f"{base_filename}.pdf"
                            pdf_path = os.path.join(pdf_dir, pdf_filename)
                            with open(pdf_path, 'wb') as f:
                                for chunk in uploaded_file.chunks():
                                    f.write(chunk)
                            document.pdf_file = os.path.join(upload_to_documents_pdf(document), pdf_filename)
                            document.save()

                            # Save PDF file to File
                            public_pdf_file = File(
                                tenant=request.tenant,
                                original_name=pdf_filename,
                                file=os.path.join(upload_to_documents_pdf(document), pdf_filename),
                                folder=template_folder,
                                uploaded_by=user
                            )
                            public_pdf_file.save()

                            print("Sending email for uploaded PDF")
                            send_approval_request(document)
                            continue

                    pdf_filename = f"{base_filename}.pdf"
                    relative_pdf_path = os.path.join(upload_to_documents_pdf(document), pdf_filename)
                    absolute_pdf_path = os.path.join(settings.MEDIA_ROOT, relative_pdf_path)

                    # Choose the right LibreOffice path
                    if platform.system() == "Windows":
                        paths = [
                            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                            r"C:\Program Files\LibreOffice\program\soffice.exe"
                        ]
                        libreoffice_path = next((p for p in paths if os.path.exists(p)), None)
                    else:
                        libreoffice_path = shutil.which("libreoffice")

                    # Check if LibreOffice exists
                    if not libreoffice_path or not os.path.exists(libreoffice_path):
                        raise FileNotFoundError("LibreOffice not found. Make sure it's installed and in PATH.")

                    try:
                        print("Starting PDF conversion with LibreOffice")
                        
                        # Ensure paths are absolute
                        abs_word_path = os.path.abspath(word_path)
                        abs_output_dir = os.path.dirname(os.path.abspath(absolute_pdf_path))

                        print("abs_word_path: ", abs_word_path)

                        # Run LibreOffice to convert .docx to .pdf
                        result = subprocess.run(
                            [libreoffice_path, "--headless", "--convert-to", "pdf", "--outdir", abs_output_dir, abs_word_path],
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            check=True
                        )

                        print("LibreOffice stdout:", result.stdout.decode())
                        print("LibreOffice stderr:", result.stderr.decode())

                        # Confirm output PDF file path
                        if os.path.exists(absolute_pdf_path):
                            document.pdf_file = relative_pdf_path
                            document.save()

                            # Save PDF file to File
                            public_pdf_file = File(
                                tenant=request.tenant,
                                original_name=pdf_filename,
                                file=relative_pdf_path,
                                folder=template_folder,
                                uploaded_by=user
                            )
                            public_pdf_file.save()
                        else:
                            return HttpResponse("PDF file was not generated.", status=500)

                    except subprocess.CalledProcessError as e:
                        print(f"LibreOffice conversion error: {e.stderr.decode()}")
                        return HttpResponse(f"Error converting to PDF: {e.stderr.decode()}", status=500)

                    except Exception as e:
                        print(f"Unexpected error: {e}")
                        return HttpResponse(f"Unexpected error converting to PDF: {e}", status=500)

                    print("Sending email")
                    send_approval_request(document)

            print("Redirecting to document_list")
            return redirect("document_list")
        else:
            print("Formset errors:", formset.errors)
            print("Non-form errors:", formset.non_form_errors())
    else:
        print("GET request received")
        formset = DocumentFormSet()
    
    return render(request, "documents/create_document.html", {"formset": formset})


def account_activation_sent(request):
    return render(request, "registration/account_activation_sent.html")

def register(request):
    if request.method == "POST":
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data["password"])
            # user.is_active = False
            user.is_active = True
            user.tenant = request.tenant
            if not request.tenant:
                return HttpResponseForbidden("No company associated with this request.")
            user.save()
            try:
                created = StaffProfile.objects.create(tenant=user.tenant, user=user)
                created.first_name = form.cleaned_data["first_name"]
                created.last_name = form.cleaned_data["last_name"]
                created.email = form.cleaned_data["email"]
                created.save()
            except ValidationError as e:
                print(f"Staff Profile creation error: {e}")
                return redirect('view_my_profile')

            # Send confirmation email
            admin_user = CustomUser.objects.filter(
                tenant=request.tenant, roles__name="Admin"
            ).first()
            if admin_user.email_address and admin_user.email_password:
                sender_email = admin_user.email_address
                sender_password = admin_user.email_password
            else:
                superuser = CustomUser.objects.get(is_superuser=True)
                sender_provider = superuser.email_provider
                sender_email = superuser.email_address
                sender_password = superuser.email_password
            if sender_email and sender_password:
                connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)
                base_domain = "127.0.0.1:8000" if settings.DEBUG else "teammanager.ng"
                protocol = "http" if settings.DEBUG else "https"
                login_url = f"{protocol}://{request.tenant.slug}.{base_domain}/accounts/login"
                subject = f"Account Pending Approval: {user.username}"
                message = f"""
                Dear {user.username},

                Your account has been created and is pending approval. You will be notified once approved. 
                Once activated, you can log in at: {login_url}

                Best regards,  
                {admin_user.get_full_name() or admin_user.username}
                """
                try:
                    send_mail(subject, message, sender_email, [user.email], connection=connection)
                except Exception as e:
                    print(f"Failed to send email: {e}")
                    # Log error or notify admin, but proceed with registration
            return redirect("account_activation_sent")
    else:
        form = SignUpForm()
    return render(request, "registration/register.html", {"form": form})

@login_required
@user_passes_test(is_admin)
def delete_user(request, user_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the user, ensuring they belong to the same tenant
    user = get_object_or_404(CustomUser, id=user_id, tenant=request.tenant)

    # Prevent admins from deleting themselves
    if user == request.user:
        return HttpResponseForbidden("You cannot delete your own account.")

    user.delete()
    return redirect("users_list")

def post_login_redirect(request):
    if not request.user.is_authenticated or request.user.is_superuser:
        return redirect('tenant_home')
    return redirect('home')

@login_required
def document_list(request):
    # Validate that the user belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: User does not belong to this company.")
    
    # if request.user.tenant.slug not in ["raadaa", "transnet-cloud"]:
    #     return HttpResponseForbidden("Unauthorized: User can not view this page.")

    # Start with documents filtered by the current tenant
    documents = Document.objects.filter(tenant=request.tenant)

    # Get filter parameters from the request
    company = request.GET.get('company', '').strip()
    doc_type = request.GET.get('type', '').strip()
    status = request.GET.get('status', '').strip()
    created = request.GET.get('created', '').strip()
    created_by = request.GET.get('created_by', '').strip()
    approved_by = request.GET.get('approved_by', '').strip()
    send_email = request.GET.get('send_email', '').strip()

    # Apply filters
    if company:
        documents = documents.filter(company_name__iexact=company)
        print(f"Filtering by company: {company}")
    if doc_type:
        documents = documents.filter(document_type__iexact=doc_type)
        print(f"Filtering by document_type: {doc_type}")
    if status:
        documents = documents.filter(status__iexact=status)
        print(f"Filtering by status: {status}")
    if created:
        try:
            documents = documents.filter(created_at__date=created)
            print(f"Filtering by created_at: {created}")
        except ValueError:
            print(f"Invalid date format for created: {created}")
    if created_by:
        documents = documents.filter(created_by__username__iexact=created_by, created_by__tenant=request.tenant)
        print(f"Filtering by created_by: {created_by}")
    if approved_by:
        documents = documents.filter(approved_by__username__iexact=approved_by, approved_by__tenant=request.tenant)
        print(f"Filtering by approved_by: {approved_by}")
    if send_email:
        email_sent = send_email.lower() == 'sent'
        documents = documents.filter(email_sent=email_sent)
        print(f"Filtering by email_sent: {email_sent}")

    # Paginate filtered documents
    paginator = Paginator(documents, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Get distinct values for filter dropdowns, scoped to the tenant
    distinct_companies = Document.objects.filter(tenant=request.tenant).values_list('company_name', flat=True).distinct()
    distinct_type = Document.objects.filter(tenant=request.tenant).values_list('document_type', flat=True).distinct()
    distinct_created_by = CustomUser.objects.filter(
        tenant=request.tenant,
        id__in=Document.objects.filter(tenant=request.tenant).values_list('created_by', flat=True).distinct()
    ).values_list('username', flat=True)
    distinct_approved_by = CustomUser.objects.filter(
        tenant=request.tenant,
        id__in=Document.objects.filter(tenant=request.tenant).values_list('approved_by', flat=True).distinct()
    ).exclude(username__isnull=True).values_list('username', flat=True)

    # Debug filtered document count
    print(f"Filtered documents count: {documents.count()}")

    return render(request, "documents/document_list.html", {
        "documents": page_obj,
        "page_obj": page_obj,
        "filter_params": {
            'company': company,
            'type': doc_type,
            'status': status,
            'created': created,
            'created_by': created_by,
            'approved_by': approved_by,
            'send_email': send_email,
        },
        "distinct_companies": distinct_companies,
        "distinct_type": distinct_type,
        "distinct_created_by": distinct_created_by,
        "distinct_approved_by": distinct_approved_by
    })

def home(request):
    tenant = request.tenant
    print("This is home")
    return render(request, "home.html", {'tenant': tenant})
   

@login_required
def approve_document(request, document_id):
    # Ensure the user belongs to the current tenant
    if not hasattr(request, 'tenant') or not request.user.tenant == request.tenant:
        return HttpResponseForbidden("You are not authorized to perform actions for this company.")

    # Fetch the document, ensuring it belongs to the current tenant
    document = get_object_or_404(Document, id=document_id, tenant=request.tenant)

    # Restrict approval to users with the BDM role
    if not request.user.roles.filter(name="BDM").exists():
        return HttpResponseForbidden("You are not allowed to approve this document.")

    # Update document status and approved_by
    document.status = "approved"
    document.approved_by = request.user
    document.save()

    # Ensure the BDM has SMTP credentials (or use tenant-specific SMTP settings)
    sender_provider = request.user.email_provider
    sender_email = request.user.email_address
    sender_password = request.user.email_password

    if not sender_email or not sender_password:
        return HttpResponseForbidden("Your email credentials are missing. Contact admin.")

    # Configure SMTP settings dynamically
    connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)

    # Ensure the document's creator belongs to the same tenant
    if document.created_by.tenant != request.tenant:
        return HttpResponseForbidden("Invalid document creator for this company.")

    # Send email notification to the BDA
    subject = f"Document Approved for {request.tenant.name}"
    message = (
        f"Hello {document.created_by.username},\n\n"
        f"Your document '{document.company_name} _ {document.document_type}' "
        f"has been approved by {request.user.username} for tenant {request.tenant.name}."
    )

    email = EmailMessage(subject, message, sender_email, [document.created_by.email], connection=connection)
    email.attach_file(document.pdf_file.path)  # Attach the PDF
    email.send()

    return redirect("document_list")  # Redirect to tenant-scoped document list

@login_required
def autocomplete_sales_rep(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    if 'term' in request.GET:
        qs = CustomUser.objects.filter(
            tenant = request.tenant,
            roles__name='Sales Rep',
            username__icontains=request.GET.get('term')
        ).distinct()
        names = list(qs.values_list('username', flat=True))
        return JsonResponse(names, safe=False)


def send_approved_email(request, document_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    document = get_object_or_404(Document, id=document_id, tenant=request.tenant)

    # Ensure the document is approved before sending an email
    if document.status != "approved":
        return HttpResponseForbidden("This document is not approved yet.")

    # Check if the email was already sent
    if document.email_sent:
        return HttpResponseForbidden("Email has already been sent.")

    # Ensure the PDF file exists
    if not document.pdf_file or not os.path.exists(document.pdf_file.path):
        return HttpResponseForbidden("PDF file not found.")

    # Get sender credentials from the logged-in 
    sender_provider = request.user.email_provider
    sender_email = request.user.email_address
    sender_password = request.user.email_password

    if not sender_email or not sender_password:
        return HttpResponseForbidden("Your email credentials are missing. Contact admin.")
    
    # Define recipient, CC, and attachment

    # Main recipient
    raw_recipients = document.contact_person_email  # e.g., "email1@example.com, email2@example.com"
    recipient = [email.strip() for email in raw_recipients.split(",") if email.strip()] # in case of multiple emails
    if not recipient:
        return HttpResponseForbidden("No valid recipient email provided.")

    # recipient = [document.contact_person_email]

    # Get Sales Rep email
    sales_rep_email = CustomUser.objects.filter(username=document.sales_rep, tenant=request.tenant).values_list("email", flat=True)

    # Get all BDM emails
    bdm_emails = CustomUser.objects.filter(roles__name="BDM", tenant=request.tenant).values_list("email", flat=True)

    # Combine into a single list
    cc_list = list(sales_rep_email) + list(bdm_emails)


    # Configure SMTP settings dynamically
    connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)

    # Email subject based on document type
    subject = (
        f"{document.company_name} - Approved by AWS"
        if document.document_type == "approval"
        else f"{document.company_name} - SLA"
    )

    # Email body based on document type
    if document.document_type == "approval":
        message = f"""
        Dear Sir,

        Trust this email finds you well.

        We are pleased to inform you that your projects have been officially approved by Amazon Web Services - AWS, under AWS Cloud Growth Services.

        Congratulations on this accomplishment which shows your business has great potential to scale and thrive on AWS platform.  
        This is a great achievement, and we are excited for you to proceed to the next phase.  
        Please find the attached document for your reference which contains the relevant details for the next steps.

        If you have any questions or need further clarification, please feel free to reach out to me.  
        Once again, congratulations and we look forward to the continued success of your projects.

        Thank you.

        Best Regards,  
        {document.created_by.get_full_name()} | Executive Assistant  
        Transnet Cloud  
        Mob: {document.created_by.phone_number}  
        No 35 Ajose Adeogun Street Utako, Abuja  
        Email: {document.created_by.email}  
        Website: www.transnetcloud.com
        """
    else:  # SLA Document
        message = f"""
        Dear {document.company_name},

        Trust this email finds you well.

        Thank you for availing us your time and attention during our last meeting.  
        We are delighted to move forward with your project on AWS to the next phase, be rest assured we will provide you with the necessary support your business needs to scale.

        Please find attached the Service Level Agreement (SLA) document for your review. Kindly take a moment to go through the terms, and if they align with your expectations, we would appreciate you signing and returning the document at your earliest convenience.

        Should you have any questions or require further clarification, please do not hesitate to reach out. Our team is readily available to assist you.

        Thank you for choosing Transnet Cloud as your trusted partner.  

        We look forward to a successful partnership.

        Best Regards,  
        {document.created_by.get_full_name()} | Executive Assistant  
        Transnet Cloud  
        Mob: {document.created_by.phone_number}  
        No 35 Ajose Adeogun Street Utako, Abuja  
        Email: {document.created_by.email}  
        Website: www.transnetcloud.com
        """

    # Create email with attachment
    email = EmailMessage(subject, message, sender_email, recipient, cc=cc_list, connection=connection)
    email.attach_file(document.pdf_file.path)  # Attach the PDF
    email.send()

    # Update email status
    document.email_sent = True
    document.save()

    return redirect("document_list")  # Redirect to the document list


# @user_passes_test(is_admin)
# def admin_access_page(request):
#     return render(request, "admin_access.html")

@login_required
@user_passes_test(is_admin)
def delete_document(request, document_id):
    if hasattr(request, 'tenant') and request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized to perform actions for this company.")
    document = get_object_or_404(Document, id=document_id, tenant=request.tenant)

    # Ensure document files are deleted from storage
    if document.word_file:
        document.word_file.delete(save=False)
    if document.pdf_file:
        document.pdf_file.delete(save=False)

    document.delete()
    return redirect("document_list")  # Redirect back to the list

def add_formatted_content(doc, soup, word_dir):
    """Parse HTML and add formatted content and images to the .docx document."""
    current_paragraph = None

    def add_text_to_paragraph(text, bold=False, italic=False):
        """Helper to add text to the current paragraph with formatting."""
        nonlocal current_paragraph
        if not current_paragraph:
            current_paragraph = doc.add_paragraph()
        run = current_paragraph.add_run(text or '')
        run.bold = bold
        run.italic = italic

    for element in soup.find_all(recursive=False):  # Process top-level elements only
        if element.name == 'h1':
            current_paragraph = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            current_paragraph = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'p':
            current_paragraph = doc.add_paragraph()
            # Process nested elements within <p>
            for child in element.children:
                if child.name == 'strong':
                    add_text_to_paragraph(child.get_text(), bold=True)
                elif child.name == 'em':
                    add_text_to_paragraph(child.get_text(), italic=True)
                elif child.name == 'img':
                    img_src = child.get('src')
                    if img_src:
                        try:
                            parsed_src = urllib.parse.urlparse(img_src)
                            if parsed_src.scheme in ('http', 'https'):
                                print(f"Downloading remote image: {img_src}")
                                response = requests.get(img_src, timeout=5)
                                if response.status_code == 200:
                                    img_data = io.BytesIO(response.content)
                                    doc.add_picture(img_data, width=Inches(4.0))
                                else:
                                    print(f"Failed to download image: {img_src} (Status: {response.status_code})")
                                    current_paragraph = doc.add_paragraph(f"[Image failed to load: {img_src}]")
                            else:
                                # Handle local images (e.g., /media/Uploads/image.jpg)
                                clean_src = img_src.replace(settings.MEDIA_URL, '').lstrip('/')
                                img_path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, clean_src))
                                print(f"Attempting to add local image: {img_path}")
                                if os.path.exists(img_path):
                                    doc.add_picture(img_path, width=Inches(4.0))
                                else:
                                    print(f"Local image not found: {img_path}")
                                    current_paragraph = doc.add_paragraph(f"[Image not found: {img_path}]")
                        except Exception as e:
                            print(f"Error adding image {img_src}: {e}")
                            current_paragraph = doc.add_paragraph(f"[Error loading image: {img_src}]")
                else:
                    add_text_to_paragraph(child.get_text() if hasattr(child, 'get_text') else str(child))
        elif element.name == 'ul':
            current_paragraph = None
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='ListBullet')
        elif element.name == 'ol':
            current_paragraph = None
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='ListNumber')
        else:
            current_paragraph = None
            add_text_to_paragraph(element.get_text() if hasattr(element, 'get_text') else str(element))

@login_required
@csrf_exempt  # Required for CKEditorâ€™s POST uploads
def custom_ckeditor_upload(request):
    if hasattr(request, 'tenant') and request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized to perform actions for this company.")
    if not request.user.is_authenticated:
        return HttpResponseForbidden("You must be logged in to upload images.")
    logger.info(f"User {request.user.username} uploading image to CKEditor")
    response = ckeditor_upload(request)
    if response.status_code == 200:
        logger.info(f"Image upload successful for user {request.user.username}")
    else:
        logger.error(f"Image upload failed for user {request.user.username}: {response.content}")
    return response

@login_required
def create_from_editor(request):
    # Validate that the user belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: User does not belong to this company.")
    
    def documents_word_upload(request):
        tenant_name = request.tenant.name
        username = request.user.username if request.user else "anonymous"
        return os.path.join('documents', tenant_name, username, 'word')

    def documents_pdf_upload(request):
        tenant_name = request.tenant.name
        username = request.user.username if request.user else "anonymous"
        return os.path.join('documents', tenant_name, username, 'pdf')

    if request.method == "POST":
        form = CreateDocumentForm(request.POST)
        if form.is_valid():
            title = form.cleaned_data["title"]
            content = form.cleaned_data["content"]

            # Create a .docx file
            doc = DocxDocument()
            doc.add_heading(title, level=1)

            # Parse HTML content using BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')

            # Define file paths
            word_dir = os.path.join(settings.MEDIA_ROOT, documents_word_upload(request))
            pdf_dir = os.path.join(settings.MEDIA_ROOT, documents_pdf_upload(request))
            os.makedirs(word_dir, exist_ok=True)
            os.makedirs(pdf_dir, exist_ok=True)

            # Add formatted content and images
            add_formatted_content(doc, soup, word_dir)

            # Get or create the "Template Document" folder
            user = request.user
            department = user.department if hasattr(user, 'department') else None
            team = user.teams.first() if hasattr(user, 'teams') and user.teams.exists() else None

            # Ensure user has a department or team
            if not department and not team:
                messages.error(request, "User must be associated with a department or team.")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Choose department or team for the folder (prefer department if available)
            folder_defaults = {
                'created_by': user,
                'is_public': True,
            }

            try:
                template_folder, created = Folder.objects.get_or_create(
                    tenant=request.tenant,
                    name="Template Document",
                    defaults=folder_defaults
                )
            except ValidationError as e:
                print(f"Folder creation error: {e}")
                messages.error(request, f"Error creating Template Document folder: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Validate folder access
            # if template_folder.department and template_folder.department != user.department:
            #     messages.error(request, "Invalid Department for Template Document folder.")
            #     return render(request, 'documents/create_from_editor.html', {'form': form})
            # if template_folder.team and template_folder.team not in user.teams.all():
            #     messages.error(request, "Invalid Team for Template Document folder.")
            #     return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save the .docx file
            word_filename = f"{slugify(title)}_{request.user.id}_{template_folder.id}.docx"
            word_path = os.path.join(word_dir, word_filename)
            try:
                doc.save(word_path)
                print(f"Saved .docx: {word_path}")
            except Exception as e:
                messages.error(request, f"Error creating .docx file: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Generate and save the .pdf file using LibreOffice
            pdf_filename = f"{slugify(title)}_{request.user.id}_{template_folder.id}.pdf"
            relative_pdf_path = os.path.join(documents_pdf_upload(request), pdf_filename)
            absolute_pdf_path = os.path.join(settings.MEDIA_ROOT, relative_pdf_path)

            # Choose the right LibreOffice path
            if platform.system() == "Windows":
                paths = [
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                    r"C:\Program Files\LibreOffice\program\soffice.exe"
                ]
                libreoffice_path = next((p for p in paths if os.path.exists(p)), None)
            else:
                libreoffice_path = shutil.which("libreoffice") or shutil.which("soffice")

            # Check if LibreOffice exists
            if not libreoffice_path or not os.path.exists(libreoffice_path):
                messages.error(request, "LibreOffice not found. Make sure it's installed and in PATH.")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            try:
                print("Starting PDF conversion with LibreOffice")
                abs_word_path = os.path.abspath(word_path)
                abs_output_dir = os.path.dirname(os.path.abspath(absolute_pdf_path))

                # Run LibreOffice to convert .docx to .pdf
                result = subprocess.run(
                    [libreoffice_path, "--headless", "--convert-to", "pdf", "--outdir", abs_output_dir, abs_word_path],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    check=True,
                    timeout=30
                )

                print("LibreOffice path:", libreoffice_path)
                print("abs_word_path:", abs_word_path)
                print("abs_output_dir:", abs_output_dir)
                print("LibreOffice stdout:", result.stdout.decode())
                print("LibreOffice stderr:", result.stderr.decode())

                # Confirm output PDF file exists
                if not os.path.exists(absolute_pdf_path):
                    if os.path.exists(word_path):
                        os.remove(word_path)
                    messages.error(request, "PDF file was not generated.")
                    return render(request, 'documents/create_from_editor.html', {'form': form})

            except subprocess.CalledProcessError as e:
                print(f"LibreOffice conversion error: {e.stderr.decode()}")
                messages.error(request, f"Error converting to PDF: {e.stderr.decode()}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            except Exception as e:
                print(f"Unexpected error during PDF conversion: {e}")
                messages.error(request, f"Unexpected error converting to PDF: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save to File
            try:
                public_word_file = File(
                    tenant=request.tenant,
                    original_name=word_filename,
                    file=os.path.join(documents_word_upload(request), word_filename),
                    folder=template_folder,
                    uploaded_by=user
                )
                public_word_file.save()

                public_pdf_file = File(
                    tenant=request.tenant,
                    original_name=pdf_filename,
                    file=relative_pdf_path,
                    folder=template_folder,
                    uploaded_by=user
                )
                public_pdf_file.save()
            except Exception as e:
                print(f"Error saving to File: {e}")
                messages.error(request, f"Error saving files to public storage: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save to Document
            document = Document(
                document_type='Uploaded',
                document_source='editor',
                company_name='N/A',
                company_address='N/A',
                contact_person_name='N/A',
                contact_person_email='N/A',
                contact_person_designation='N/A',
                sales_rep='N/A',
                created_by=request.user,
                tenant=request.tenant,
                word_file=os.path.join(documents_pdf_upload(request), word_filename),
                pdf_file=relative_pdf_path
            )
            try:
                document.save()
            except Exception as e:
                messages.error(request, f"Error saving document: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            messages.success(request, 'Document created successfully!')
            return redirect("document_list")
        else:
            print("Form errors:", form.errors)
            messages.error(request, 'Please correct the errors below.')
    else:
        form = CreateDocumentForm()
    return render(request, 'documents/create_from_editor.html', {'form': form})

@login_required
def folder_view(request, public_folder_id=None, personal_folder_id=None):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")

    # Get active tab, default to 'public'
    active_tab = request.GET.get('tab', 'public')

    # Public tab context
    public_parent = None
    if public_folder_id:
        public_parent = get_object_or_404(
            Folder,
            id=public_folder_id,
            # created_by=request.user,
            tenant=request.tenant,
            is_public=True
        )
    public_folders = Folder.objects.filter(
        # created_by=request.user,
        parent=public_parent,
        tenant=request.tenant,
        is_public=True
    )
    public_files = File.objects.filter(
        folder=public_parent,
        # uploaded_by=request.user,
        tenant=request.tenant
    )

    # Personal tab context
    personal_parent = None
    if personal_folder_id:
        personal_parent = get_object_or_404(
            Folder,
            id=personal_folder_id,
            created_by=request.user,
            tenant=request.tenant,
            is_public=False
        )
    personal_folders = Folder.objects.filter(
        created_by=request.user,
        parent=personal_parent,
        tenant=request.tenant,
        is_public=False
    )
    personal_files_1 = File.objects.filter(
        folder=personal_parent,
        uploaded_by=request.user,
        tenant=request.tenant,
        is_public = False
    )
    personal_files_2 = File.objects.filter(
        folder=personal_parent,
        # uploaded_by=request.user,
        tenant=request.tenant,
        is_public=False
    )
    personal_files = personal_files_1.union(personal_files_2)

    now = timezone.now()
    # today = now.date()

    # Generate shareable links for files
    for file in public_files:
        file.shareable_link = request.build_absolute_uri(file.get_shareable_link()) if file.is_shared else None
    for file in personal_files:
        file.shareable_link = request.build_absolute_uri(file.get_shareable_link()) if file.is_shared else None

    # Generate shareable links for folders
    for folder in public_folders:
        folder.shareable_link = request.build_absolute_uri(folder.get_shareable_link()) if folder.is_shared else None
        if folder.share_time_end and folder.share_time_end < now:
            folder.is_shared = False
            # folder.save()
            # folder.shareable_link = None
    for folder in personal_folders:
        folder.shareable_link = request.build_absolute_uri(folder.get_shareable_link()) if folder.is_shared else None
        if folder.share_time_end and folder.share_time_end < now:
            folder.is_shared = False
            # folder.save()
            # folder.shareable_link = None


    folder_form = FolderForm(initial={'parent': public_parent if active_tab == 'public' else personal_parent})
    file_form = FileUploadForm()

    return render(request, 'folder/folders.html', {
        'active_tab': active_tab,
        'public_parent': public_parent,
        'public_folders': public_folders,
        'public_files': public_files,
        'personal_parent': personal_parent,
        'personal_folders': personal_folders,
        'personal_files': personal_files,
        'folder_form': folder_form,
        'file_form': file_form,
    })

@login_required
def create_folder(request):
    if request.method == 'POST':
        form = FolderForm(request.POST)
        active_tab = request.POST.get('tab')  # Changed from request.GET to request.POST
        if form.is_valid():
            folder = form.save(commit=False)
            folder.created_by = request.user
            folder.tenant = request.tenant
            
            if active_tab == 'public':
                folder.is_public = True
            else:
                folder.is_public = False
                
            # Validate that the user belongs to the same tenant
            if folder.created_by.tenant != request.tenant:
                return JsonResponse({'success': False, 'errors': 'Unauthorized: User does not belong to this company.'}, status=403)
            
            # Get parent ID from POST data, not GET
            parent_id = request.POST.get('parent')
            print(f"Parent ID: {parent_id}, Active tab: {active_tab}")
            
            if parent_id and parent_id != 'None' and parent_id != '':
                try:
                    folder.parent = Folder.objects.get(id=parent_id, tenant=request.tenant)
                except Folder.DoesNotExist:
                    return JsonResponse({'success': False, 'errors': 'Parent folder not found.'}, status=400)
            
            folder.save()
            return JsonResponse({'success': True, 'redirect_url': f'/folders/?tab={active_tab}'})
        else:
            return JsonResponse({'success': False, 'errors': form.errors}, status=400)
    
    return JsonResponse({'success': False, 'errors': 'Invalid request method.'}, status=400)

@login_required
def delete_folder(request, folder_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    folder = get_object_or_404(Folder, id=folder_id, created_by=request.user, tenant=request.tenant)
    if folder.created_by != request.user:
        return HttpResponseForbidden("You are not authorized to delete this folder.")
    folder.delete()
    return JsonResponse({'success': True})
    # return JsonResponse({'success': False, 'errors': form.errors}, status=400)

@login_required
def upload_file(request, public_folder_id=None, personal_folder_id=None):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return JsonResponse({"success": False, "errors": "You are not authorized for this company."}, status=403)

    active_tab = request.POST.get('tab')  # Changed from request.GET to request.POST
    
    if request.method == 'POST':
        form = FileUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.save(commit=False)
            uploaded_file.uploaded_by = request.user
            uploaded_file.tenant = request.tenant
            uploaded_file.original_name = request.FILES['file'].name
            
            # Get folder from POST data or URL parameters
            folder_id = request.POST.get('folder')
            if folder_id and folder_id != 'None' and folder_id != '':
                try:
                    uploaded_file.folder = Folder.objects.get(id=folder_id, tenant=request.tenant)
                    uploaded_file.is_public = uploaded_file.folder.is_public
                except Folder.DoesNotExist:
                    return JsonResponse({"success": False, "errors": "Folder not found."}, status=400)
            
            uploaded_file.save()
            return JsonResponse({"success": True, "redirect_url": f'/folders/?tab={active_tab}'})
        else:
            return JsonResponse({"success": False, "errors": form.errors}, status=400)
    
    return JsonResponse({"success": False, "errors": "Invalid request method."}, status=400)

from documents.forms import FileUploadAnonForm
def upload_file_anon(request, public_folder_id=None, personal_folder_id=None):
    # Check tenant authorization first
    # if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
    #     return HttpResponseForbidden("You are not authorized for this tenant.")

    # active_tab = request.GET.get('tab', 'public')
    print("Public folder id:", public_folder_id)
    public_parent_folder = None
    personal_parent_folder = None
    if public_folder_id:
        public_parent_folder = Folder.objects.get(
            id=public_folder_id, 
            tenant=request.tenant, 
            is_public=True
        )
    elif personal_folder_id:
        personal_parent_folder = Folder.objects.get(
            id=personal_folder_id,
            tenant=request.tenant, 
            is_public=False
        )
    else:
        folder_id = None

    if request.method == 'POST':
        form = FileUploadAnonForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = form.save(commit=False)
            # uploaded_file.uploaded_by = None
            uploaded_file.anon_name = form.cleaned_data["anon_name"]
            # print("uploaded by", uploaded_file.uploaded_by.username)
            print("uploaded by: ", form.cleaned_data["anon_name"])
            uploaded_file.tenant = request.tenant
            uploaded_file.original_name = request.FILES['file'].name
            if public_parent_folder:
                uploaded_file.folder = public_parent_folder  # Assign to the current folder
            else:
                uploaded_file.folder = personal_parent_folder
            if uploaded_file.folder.is_public:
                uploaded_file.is_public = True
            uploaded_file.save()
            
            return JsonResponse({"success": True, "file_id": uploaded_file.id})
        else:
            # If form is invalid, re-render the form with errors
            # Re-render the form with errors (adjust context as needed)
            return render(request, 'folder/folders.html', {
                'file_form': form,
                'public_parent': public_parent_folder,
                'personal_parent': personal_parent_folder,
                'public_folders': Folder.objects.filter(created_by=request.user, parent=public_parent_folder, tenant=request.tenant, is_public=True),
                'public_files': File.objects.filter(folder=public_parent_folder, uploaded_by=request.user, tenant=request.tenant),
                'personal_folders': Folder.objects.filter(created_by=request.user, parent=personal_parent_folder, tenant=request.tenant, is_public=False),
                'personal_files': File.objects.filter(folder=personal_parent_folder, uploaded_by=request.user, tenant=request.tenant),
                'folder_form': FolderForm(initial={'parent': public_parent_folder if public_folder_id else personal_parent_folder}),
            })
    else:
        return JsonResponse({"success": False, "errors": form.errors})
    
@login_required
def enable_folder_sharing(request, folder_id):
    # View to toggle the sharing status of a file
    folder = get_object_or_404(Folder, id=folder_id, tenant=request.user.tenant)
    
    # Optional: Add permission checks (e.g., only allow certain users to toggle sharing)
    # if public_file.created_by != request.user:
    #     # return HttpResponseForbidden("You do not have permission to modify this file.")
    #     raise PermissionDenied("You do not have permission to modify this file.")

    if request.method == 'POST':
        end_date = request.POST.get('end_date')  # Already handled; could be empty string or None
        share_subfolders = 'share_folders' in request.POST  # True if checked, False otherwise
        share_files = 'share_files' in request.POST  # True if checked, False otherwise

        folder.is_shared = not folder.is_shared
        folder.share_time = timezone.now()
        folder.share_time_end = end_date if end_date else None
        folder.shared_by = request.user
        folder.share_subfolders = share_subfolders
        folder.share_files = share_files

        folder.save()
        if end_date:
            folder.share_time_end = end_date
        if share_subfolders:
            folder.share_subfolders = True
            for subfolder in Folder.objects.filter(parent=folder, tenant=request.tenant):
                subfolder.is_shared = True
                subfolder.share_time = timezone.now()
                subfolder.share_time_end = end_date if end_date else None
                subfolder.shared_by = request.user
                subfolder.share_subfolders = share_subfolders
                subfolder.share_files = share_files
                subfolder.save()
        if share_files:
            for file in folder.files.all():  # Adjust based on your model
                    file.is_shared = True
                    file.share_time = timezone.now()
                    file.share_time_end = end_date if end_date else None
                    file.shared_by = request.user
                    file.save()
        # folder.shared_by = request.user
        # folder.save()
        active_tab = request.GET.get('tab', 'public')
        
    # Determine the correct URL based on folder.is_public and active_tab
    if folder.parent:
        if folder.is_public:
            url_name = 'folder_view_public'
            kwargs = {'public_folder_id': folder.parent.id}
        else:
            url_name = 'folder_view_personal'
            kwargs = {'personal_folder_id': folder.parent.id}
    else:
        url_name = 'folder_view'
        kwargs = {}

    # return redirect(f"{reverse(url_name, kwargs=kwargs)}?tab={active_tab}")
    return JsonResponse({"success": True, "folder_id": folder.id})

@login_required
def enable_file_sharing(request, file_id):
    # View to toggle the sharing status of a file
    file = get_object_or_404(File, id=file_id, tenant=request.user.tenant)
    end_date = request.POST.get('end_date')

    if request.method == 'POST':
        file.is_shared = not file.is_shared
        file.share_time = timezone.now()
        if end_date:
            file.share_time_end = end_date
        file.shared_by = request.user
        file.save()
        active_tab = request.GET.get('tab', 'public')
        
    # Determine the correct URL based on file.is_public and active_tab
    if file.folder:
        if file.is_public:
            url_name = 'folder_view_public'
            kwargs = {'public_folder_id': file.folder.id}
        else:
            url_name = 'folder_view_personal'
        kwargs = {'personal_folder_id': file.folder.id}
    else:
        url_name = 'folder_view'
        kwargs = {}

    # return redirect(f"{reverse(url_name, kwargs=kwargs)}?tab={active_tab}")
    return JsonResponse({"success": True, "file_id": file.id})

def shared_folder_view(request, token):
    # Retrieve the folder by share token
    folder = get_object_or_404(Folder, share_token=token, is_shared=True)
    folders = Folder.objects.filter(parent=folder, tenant=request.tenant)
    files = File.objects.filter(
        folder=folder,
        # uploaded_by=request.user,
        tenant=request.tenant
    )
    file_form = FileUploadAnonForm()
    
    # Optional: Add additional checks (e.g., tenant status, file availability)
    if not folder.name:
        return HttpResponseForbidden("File not available.")
    
    for fold in folders:
        fold.is_shared = True
        fold.shareable_link = request.build_absolute_uri(fold.get_shareable_link())
        fold.share_time = timezone.now()
        fold.share_time_end = fold.parent.share_time_end
        fold.shared_by = fold.parent.shared_by
        fold.save()
    
    context = {
        'folder': folder,
        'folders': folders if folder.share_subfolders else [],
        'files': files if folder.share_files else [],
        'file_form': file_form,
        # 'file_url': file.file.url,  # URL to access the file
    }
    return render(request, 'folder/shared_folder_view.html', context)

def shared_file_view(request, token):
    # Retrieve the file by share token
    file = get_object_or_404(File, share_token=token, is_shared=True)
    
    # Optional: Add additional checks (e.g., tenant status, file availability)
    if not file.file:
        return HttpResponseForbidden("File not available.")
    
    context = {
        'file': file,
        'file_url': file.file.url,  # URL to access the file
    }
    return render(request, 'folder/shared_file_view.html', context)


@login_required
def delete_file(request, file_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    file = get_object_or_404(File, id=file_id, uploaded_by=request.user, tenant=request.tenant)
    if file.uploaded_by != request.user:
        return HttpResponseForbidden("You are not authorized to delete this file.")
    file.delete()
    return JsonResponse({'success': True})
    # return JsonResponse({'success': False, 'errors': form.errors}, status=400)

@login_required
def rename_folder(request, folder_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    folder = get_object_or_404(Folder, id=folder_id, created_by=request.user, tenant=request.tenant)
    if folder.created_by != request.user:
        return HttpResponseForbidden("You are not authorized to rename this folder.")
    if request.method == 'POST':
        new_name = request.POST.get('name', '').strip()

        # Validate the new name
        if not new_name:
            return JsonResponse({'success': False, 'error': 'Folder name cannot be empty.'}, status=400)

        if len(new_name) > 255:
            return JsonResponse({'success': False, 'error': 'Folder name is too long.'}, status=400)

        # Validate name format (e.g., no special characters)
        if not re.match(r'^[\w\s\-\.]+$', new_name):
            return JsonResponse({'success': False, 'error': 'Folder name contains invalid characters.'}, status=400)

        # Check for duplicate folder names within the tenant
        if Folder.objects.filter(tenant=request.tenant, name=new_name).exclude(id=folder.id).exists():
            return JsonResponse({'success': False, 'error': 'A folder with this name already exists.'}, status=400)

        try:
            with transaction.atomic():
                folder.name = new_name
                folder.save()
                return JsonResponse({'success': True, 'new_name': folder.name})
        except ValidationError as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=400)

    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=400)


@login_required
def rename_file(request, file_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    file = get_object_or_404(File, id=file_id, uploaded_by=request.user, tenant=request.tenant)
    if request.method == 'POST':
        new_name = request.POST.get('name')
        # Validate the new name
        if not new_name:
            return JsonResponse({'success': False, 'error': 'File name cannot be empty.'}, status=400)

        if len(new_name) > 255:
            return JsonResponse({'success': False, 'error': 'File name is too long.'}, status=400)

        # Validate name format (e.g., no special characters)
        if not re.match(r'^[\w\s\-\.]+$', new_name):
            return JsonResponse({'success': False, 'error': 'File name contains invalid characters.'}, status=400)

        try:
            with transaction.atomic():
                file.original_name = new_name
                file.save()
                return JsonResponse({'success': True, 'new_name': file.original_name})
        except ValidationError as e:
            return JsonResponse({'success': False, 'error': str(e)}, status=400)

    return JsonResponse({'success': False, 'error': 'Invalid request method.'}, status=400)

@login_required
def move_folder(request, folder_id):
    if not hasattr (request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    folder = get_object_or_404(Folder, id=folder_id, created_by=request.user, tenant=request.tenant)
    if request.method == 'POST':
        active_tab = request.POST.get('tab')
        if folder.is_public != (active_tab == 'public'):
            return JsonResponse({'success': False, 'errors': {'general': 'Invalid tab context'}})
        new_parent_id = request.POST.get('new_parent_id')
        if new_parent_id:
            folder.parent = Folder.objects.get(id=new_parent_id, tenant=request.tenant)
            folder.save()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False}, status=400)


@login_required
def move_file(request, file_id):
    if not hasattr (request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")
    file = get_object_or_404(File, id=file_id, uploaded_by=request.user, tenant=request.tenant)
    if request.method == 'POST':
        active_tab = request.POST.get('tab')
        if file.is_public != (active_tab == 'public'):
            return JsonResponse({'success': False, 'errors': {'general': 'Invalid tab context'}})
        new_folder_id = request.POST.get('new_folder_id')
        if new_folder_id:
            file.folder = Folder.objects.get(id=new_folder_id, tenant=request.tenant)
            file.save()
        return JsonResponse({'success': True})
    return JsonResponse({'success': False}, status=400)

# Tasks
@login_required
def create_task(request):
    # Ensure the user belongs to the current tenant
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized for this company.")

    if request.method == 'POST':
        form = TaskForm(request.POST, user=request.user)
        if form.is_valid():
            task = form.save(commit=False)
            task.created_by = request.user
            task.tenant = request.user.tenant
            task.save()
            form.save_m2m()
            return redirect('task_list')
    else:
        form = TaskForm(user=request.user)

    return render(request, 'tasks/create_task.html', {'form': form})


@login_required
def task_list(request):
    # Validate tenant access
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")

    category = request.GET.get('category', 'overall')
    
    # Filter tasks by tenant and user (assigned_to or created_by)
    tasks = Task.objects.filter(
        Q(assigned_to=request.user) | Q(created_by=request.user),
        tenant=request.tenant
    ).distinct()

    reassign_form = ReassignTaskForm(user=request.user)
    
    # Apply category-specific filtering
    if category == 'personal':
        tasks = tasks.filter(assigned_to=request.user, created_by=request.user)
    elif category == 'corporate':
        tasks = tasks.filter(assigned_to=request.user).exclude(created_by=request.user)
    
    # Update overdue tasks for the current tenant
    overdue_tasks = Task.objects.filter(
        tenant=request.tenant,
        due_date__lt=date.today(),
        status__in=['pending', 'in_progress', 'on_hold'],
        assigned_to=request.user
    )
    if overdue_tasks.exists():
        logger.debug(f"Updating {overdue_tasks.count()} overdue tasks for tenant {request.tenant}")
        overdue_tasks.update(status='overdue')

    # Filter users by tenant and optionally by department
    users = CustomUser.objects.filter(tenant=request.tenant)
    if request.user.department:
        users = users.filter(department=request.user.department)
    
    logger.debug(f"Task list for tenant {request.tenant}: {tasks.count()} tasks, {users.count()} users")

    context = {
        'tasks': tasks,
        'status_labels': Task.STATUS_CHOICES,
        'today': date.today(),
        'user': request.user,
        'users': users,
        'category': category,
        'reassign_form': reassign_form,
    }
    return render(request, 'tasks/task_list.html', context)

@login_required
def task_detail(request, task_id):
    task = get_object_or_404(Task, id=task_id, tenant=request.tenant)
    user = request.user
    if not (user in task.assigned_to.all() or task.created_by == request.user or request.user.is_hod()):
        return render(request, 'tasks/error.html', {'message': 'Access denied.'})
    return render(request, 'tasks/task_detail.html', {'task': task})

@csrf_exempt
@login_required
def update_task_status(request, task_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.") 
    task = get_object_or_404(Task, id=task_id, tenant=request.tenant)
    if not (task.assigned_to == request.user or task.created_by == request.user or request.user.is_hod()):
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            new_status = data.get('status')
            if new_status in dict(Task.STATUS_CHOICES):
                task.status = new_status
                if new_status == 'completed':
                    task.completed_at = timezone.now()
                task.tenant = request.tenant
                task.save()
                return JsonResponse({'success': True, 'status': task.status})
            return JsonResponse({'error': 'Invalid status'}, status=400)
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required
def reassign_task(request, task_id):
    # Fetch the task, ensuring it belongs to the current tenant
    task = get_object_or_404(Task, id=task_id, tenant=request.tenant)

    # Check permissions: only task creator or HOD can reassign
    if not (task.created_by == request.user or request.user.is_hod()):
        return JsonResponse({'error': 'Access denied.'}, status=403)

    if request.method == 'POST':
        try:
            # Extract assigned user IDs from POST request
            assigned_to_ids = request.POST.getlist('assigned_to')
            due_date_str = request.POST.get('due_date')

            # Validate inputs
            if not assigned_to_ids:
                return JsonResponse({'error': 'At least one assigned user is required.'}, status=400)
            if not due_date_str:
                return JsonResponse({'error': 'Due date is required.'}, status=400)

            # Parse and validate due date
            try:
                due_date = datetime.strptime(due_date_str, '%Y-%m-%d').date()
                if due_date < date.today():
                    return JsonResponse({'error': 'Due date cannot be in the past.'}, status=400)
            except ValueError:
                return JsonResponse({'error': 'Invalid due date format. Use YYYY-MM-DD.'}, status=400)

            # Fetch assigned users, ensuring they belong to the same tenant
            assigned_users = CustomUser.objects.filter(id__in=assigned_to_ids, tenant=request.tenant)
            if not assigned_users.exists() or len(assigned_users) != len(assigned_to_ids):
                return JsonResponse({'error': 'One or more selected users are invalid or not in the tenant.'}, status=400)

            # Update task fields
            task.due_date = due_date
            task.status = 'in_progress'
            task.save()

            # Update many-to-many relationship for assigned users
            task.assigned_to.clear()  # Clear existing assignments
            task.assigned_to.add(*assigned_users)  # Add new assignments

            return JsonResponse({'success': True, 'message': 'Task reassigned successfully.'})
        except ValidationError as e:
            return JsonResponse({'error': str(e)}, status=400)
        except Exception as e:
            return JsonResponse({'error': f'An error occurred: {str(e)}'}, status=500)

    return JsonResponse({'error': 'Invalid request method. Use POST.'}, status=405)

@login_required
def delete_task(request, task_id):
    task = get_object_or_404(Task, id=task_id, tenant=request.tenant)
    if not (task.created_by == request.user or request.user.is_hod()):
        return render(request, 'tasks/error.html', {'message': 'Access denied.'})
    
    if request.method == 'POST':
        task.delete()
        return redirect('task_list')
    return render(request, 'tasks/confirm_delete.html', {'task': task})

@login_required
def task_edit(request, task_id):
    # Validate tenant access
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")

    # Fetch task with tenant filter
    task = get_object_or_404(Task, id=task_id, tenant=request.user.tenant)

    # Check access permissions
    if not (task.created_by == request.user or request.user.is_hod()):
        logger.warning(f"Access denied for user {request.user.username} on task {task_id}")
        return render(request, 'tasks/error.html', {'message': 'Access denied.'})
    
    if request.method == 'POST':
        form = TaskForm(request.POST, request.FILES, instance=task, user=request.user)
        if form.is_valid():
            form.instance.tenant = request.tenant
            form.save()
            logger.info(f"Task {task.id} edited by {request.user.username} in tenant {request.tenant}")
            return redirect('task_detail', task_id=task.id)
        
    else:
        form = TaskForm(instance=task, user=request.user)
        logger.debug(f"TaskForm initialized for task {task_id} with tenant: {request.tenant}")

    return render(request, 'tasks/edit_task.html', {'form': form, 'task': task})

@login_required
def delete_task_document(request, task_id, file_id):
    task = get_object_or_404(Task, id=task_id, tenant=request.tenant)
    document = get_object_or_404(File, id=file_id, tenant=request.tenant)
    if not (task.created_by == request.user or request.user.is_hod()) or document not in task.documents.all():
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method == 'POST':
        task.documents.remove(document)
        if not task.documents.exists():
            document.delete()
        return JsonResponse({'success': True})
    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required
def performance_dashboard(request):
    # Validate tenant access
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")

    user = request.user
    category = request.GET.get('category', 'overall')
    
    # Filter tasks by tenant and assigned_to user
    tasks = Task.objects.filter(tenant=request.user.tenant, assigned_to=user)
    if category == 'personal':
        tasks = tasks.filter(created_by=user)
    elif category == 'corporate':
        tasks = tasks.exclude(created_by=user)
    
    total_tasks = tasks.count()
    completed_tasks = tasks.filter(status='completed').count()
    completion_percentage = (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0

    now = timezone.now()
    week_ago = now - timedelta(days=7)
    month_ago = now - timedelta(days=30)
    year_ago = now - timedelta(days=365)

    weekly_completed = tasks.filter(status='completed', completed_at__gte=week_ago).count()
    monthly_completed = tasks.filter(status='completed', completed_at__gte=month_ago).count()
    yearly_completed = tasks.filter(status='completed', completed_at__gte=year_ago).count()

    overdue_tasks = tasks.filter(status='overdue', due_date__lt=now).count()

    performance_score = completion_percentage - (overdue_tasks * 10)
    performance_score = max(0, min(100, performance_score))

    context = {
        'category': category,
        'completion_percentage': round(completion_percentage, 2),
        'weekly_completed': weekly_completed,
        'monthly_completed': monthly_completed,
        'yearly_completed': yearly_completed,
        'overdue_tasks': overdue_tasks,
        'performance_score': round(performance_score, 2),
    }
    return render(request, 'dashboard/performance_dashboard.html', context)

@login_required
def hod_performance_dashboard(request):
    # Validate tenant access
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    user = request.user
    if not user.is_hod():
        return render(request, 'tasks/error.html', {'message': 'Access denied. HOD role required.'})

    # Get all departments where user is HOD
    departments = Department.objects.filter(hod=user, tenant=request.tenant)
    if not departments.exists():
        return render(request, 'tasks/error.html', {'message': 'No departments assigned to you.'})

    # Get department filter from query parameter
    selected_department_id = request.GET.get('department_id', 'all')
    
    # Filter departments if a specific one is selected
    if selected_department_id != 'all':
        try:
            selected_department = departments.get(id=selected_department_id)
            departments_to_show = [selected_department]
        except Department.DoesNotExist:
            departments_to_show = departments
            selected_department_id = 'all'
    else:
        departments_to_show = departments

    # Get all users from the selected department(s)
    department_users = CustomUser.objects.filter(
        department__in=departments_to_show, 
        tenant=request.tenant
    ).select_related('department')
    
    users_ids = [u.id for u in department_users]

    # Get user filter from query parameter (optional)
    selected_user_id = request.GET.get('user_id', 'all')
    
    # Base query: corporate tasks (created by department members, not personal)
    tasks = Task.objects.filter(
        created_by__in=users_ids, 
        tenant=request.tenant
    ).exclude(created_by=F('assigned_to')).select_related('created_by', 'assigned_to')

    # Filter by selected user if specified
    if selected_user_id != 'all':
        tasks = tasks.filter(assigned_to__id=selected_user_id)

    # Department-wide or user-specific metrics
    total_tasks = tasks.count()
    completed_tasks = tasks.filter(status='completed').count()
    completion_percentage = (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0

    # Per-user metrics for the table
    user_metrics = []
    for dept_user in department_users:
        user_tasks = Task.objects.filter(
            created_by__in=users_ids,
            assigned_to=dept_user,
            tenant=request.tenant
        ).exclude(created_by=F('assigned_to')).select_related('assigned_to')
        
        user_total_tasks = user_tasks.count()
        user_completed_tasks = user_tasks.filter(status='completed').count()
        user_completion_percentage = (user_completed_tasks / user_total_tasks * 100) if user_total_tasks > 0 else 0
        
        user_metrics.append({
            'user_id': dept_user.id,
            'full_name': dept_user.get_full_name() or dept_user.username,
            'department': dept_user.department.name if dept_user.department else 'N/A',
            'total_tasks': user_total_tasks,
            'completed_tasks': user_completed_tasks,
            'completion_percentage': round(user_completion_percentage, 2),
        })

    context = {
        'departments': departments,
        'selected_department_id': selected_department_id,
        'completion_percentage': round(completion_percentage, 2),
        'total_tasks': total_tasks,
        'completed_tasks': completed_tasks,
        'user_metrics': user_metrics,
        'department_users': department_users,
        'selected_user_id': selected_user_id,
    }
    return render(request, 'dashboard/hod_performance_dashboard.html', context)
    
@login_required
def view_my_profile(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    profile, created = StaffProfile.objects.get_or_create(user=request.user, tenant=request.tenant)
    visible_fields = [
            "photo",
            "first_name", "last_name", "middle_name", "email", "phone_number", "sex", "date_of_birth", "home_address",
            "state_of_origin", "lga", "religion",
            "institution", "course", "degree", "graduation_year",
            "account_number", "bank_name", "account_name",
            "location", "employment_date", "department", "team", "designation", "official_email",
            "emergency_name", "emergency_relationship", "emergency_phone",
            "emergency_address", "emergency_email",
        ]
    return render(request, "dashboard/my_profile.html", {"profile": profile, "visible_fields": visible_fields})


@login_required
def edit_my_profile(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    profile,_ = StaffProfile.objects.get_or_create(user=request.user, tenant=request.tenant) # staff_profile = request.user.staff_profile  # Assuming one profile per user
    if request.method == 'POST':
        profile_form = StaffProfileForm(request.POST, request.FILES, instance=profile, user=request.user)
        
        if profile_form.is_valid():
            profile_form.save()
            return redirect('view_my_profile')  # Redirect to profile view or success page
    else:
        profile_form = StaffProfileForm(instance=profile, user=request.user)
        document_form = StaffDocumentForm()
        
    return render(request, 'dashboard/edit_profile.html', {
        'profile': profile,
        'profile_form': profile_form,
        'document_form': document_form,
    })

@login_required
def add_staff_document(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    if request.method == 'POST':
        form = StaffDocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.tenant = request.user.tenant
            document.staff_profile = get_object_or_404(StaffProfile, user=request.user, tenant=request.tenant)
            document.save()
            return JsonResponse({
                'success': True,
                'document': {
                    'id': document.id,
                    'description': document.description or document.document_type,
                    'file_url': document.file.url,
                    'document_type': document.get_document_type_display(),
                    'uploaded_at': document.uploaded_at.strftime('%B %d, %Y')
                }
            })
        else:
            return JsonResponse({
                'success': False,
                'errors': form.errors
            }, status=400)
    return JsonResponse({'success': False, 'error': 'Invalid request'}, status=400)

@login_required
def delete_staff_document(request, document_id):
    if hasattr(request, 'tenant') and request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    try:
        # Get the user's StaffProfile (assuming one profile per user)
        staff_profile = StaffProfile.objects.get(user=request.user, tenant=request.tenant)
        document = get_object_or_404(StaffDocument, id=document_id, staff_profile=staff_profile)
        if request.user != document.staff_profile.user:
            return JsonResponse({'success': False, 'error': 'You are not authorized to delete this document.'}, status=403)
        if request.method == 'POST':
            document.delete()
            return JsonResponse({'success': True})
        return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)
    except StaffProfile.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Staff profile not found'}, status=404)
    except StaffDocument.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Document not found or not owned by user'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@login_required
@user_passes_test(is_admin)
def staff_directory(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    # Filter staff profiles by tenant
    profiles = StaffProfile.objects.select_related(
        "user", "department"
    ).prefetch_related("team").filter(user__tenant=request.user.tenant)

    # Restrict to department for HODs (optional)
    if request.user.is_hod() and not is_admin(request.user):
        profiles = profiles.filter(department=request.user.department)

    logger.debug(f"Profiles found: {profiles.count()}")

    # Grouping by Department â†’ Team (Tenant is already scoped)
    grouped = {}
    for profile in profiles:
        dept = profile.department.name if profile.department else "No Department"
        teams = profile.team.all()
        if not teams:
            teams = [None]
        for team in teams:
            team_name = team.name if team else "No Team"
            grouped.setdefault(dept, {}).setdefault(team_name, []).append(profile)

    logger.debug(f"Grouped dictionary: {grouped}")
    return render(request, "staff/staff_directory.html", {"grouped": grouped})


@login_required
def view_staff_profile(request, user_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    profile = get_object_or_404(StaffProfile, user_id=user_id, tenant=request.tenant)

    # Ensure the profile belongs to the same tenant as the requesting user
    if profile.user.tenant != request.user.tenant:
        return render(request, "403.html", status=403)

    viewer = StaffProfile.objects.filter(user=request.user, tenant=request.tenant).first()

    visible_fields = [
        "photo",
        "first_name", "last_name", "middle_name", "email", "phone_number", "sex", "date_of_birth", "home_address",
        "state_of_origin", "lga", "religion",
        "institution", "course", "degree", "graduation_year",
        "account_number", "bank_name", "account_name",
        "location", "employment_date",
        "department", "team", "designation", "official_email",
        "emergency_name", "emergency_relationship", "emergency_phone",
        "emergency_address", "emergency_email",
    ]

    # Visibility logic based on shared department/team (within same tenant already)
    if viewer:
        if profile.department == viewer.department:
            visible_fields += ["department"]
            shared_teams = set(profile.team.values_list("id", flat=True)) & set(viewer.team.values_list("id", flat=True))
            if shared_teams:
                visible_fields += ["team"]

    return render(request, "staff/view_staff_profile.html", {
        "profile": profile,
        "viewer": viewer,
        "visible_fields": visible_fields,
    })

def staff_list(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    # Get query parameters
    sort_by = request.GET.get('sort_by', 'name')
    sort_order = request.GET.get('sort_order', 'asc')
    search_query = request.GET.get('search', '')
    filter_dept = request.GET.get('dept', '')
    page = request.GET.get('page', 1)

    # Base queryset: only staff from the same tenant
    profiles = StaffProfile.objects.select_related(
        'user', 'department'
    ).prefetch_related('team').filter(user__tenant=request.user.tenant)

    # Search by name
    if search_query:
        profiles = profiles.filter(
            Q(first_name__icontains=search_query) |
            Q(middle_name__icontains=search_query) |
            Q(last_name__icontains=search_query)
        )

    # Filter by department
    if filter_dept:
        profiles = profiles.filter(department__name__iexact=filter_dept)

    # Sorting
    sort_field = sort_by
    if sort_by == 'name':
        sort_field = 'first_name'
    elif sort_by == 'department':
        sort_field = 'department__name'
    elif sort_by == 'team':
        sort_field = 'team__name'
    elif sort_by == 'photo':
        sort_field = 'photo'

    if sort_order == 'desc':
        sort_field = f'-{sort_field}'
    profiles = profiles.order_by(sort_field)

    # Pagination
    paginator = Paginator(profiles, 10)  # 10 profiles per page
    page_obj = paginator.get_page(page)

    # Filter dropdown: restrict departments to the tenant only
    departments = Department.objects.filter(
        staff__user__tenant=request.user.tenant
    ).distinct()

    context = {
        'profiles': page_obj,
        'departments': departments,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'search_query': search_query,
        'filter_dept': filter_dept,
    }
    return render(request, 'staff/staff_list.html', context)

from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.views.decorators.http import require_POST
from .models import Notification, UserNotification
from django.utils.timezone import now

@login_required
def notifications_view(request):
    # Validate tenant access
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")

    # Filter notifications by tenant and active status
    all_notifications = Notification.objects.filter(
        tenant=request.user.tenant,
        is_active=True
    )

    # Ensure UserNotification exists for each active notification
    active_notifications = []
    for notification in all_notifications:
        user_notification, created = UserNotification.objects.get_or_create(
            user=request.user,
            tenant=request.user.tenant,
            notification=notification,
            defaults={'seen_at': timezone.now(), 'dismissed': False}
        )
        if not user_notification.dismissed:
            active_notifications.append(user_notification)

    # Get dismissed notifications for the user and tenant
    dismissed_notifications = UserNotification.objects.filter(
        user=request.user,
        notification__tenant=request.user.tenant,
        dismissed=True
    ).select_related('notification').order_by('-seen_at')

    return render(request, 'users/notifications.html', {
        'active_notifications': active_notifications,
        'dismissed_notifications': dismissed_notifications
    })

@require_POST
@login_required
def dismiss_notification(request):
    notification_id = request.POST.get('notification_id')
    try:
        notification = Notification.objects.get(id=notification_id, tenant=request.user.tenant)
        user_notification, created = UserNotification.objects.get_or_create(
            tenant=request.user.tenant,
            user=request.user,
            notification=notification,
            defaults={'seen_at': now()}
        )
        user_notification.dismissed = True
        user_notification.save()
        return JsonResponse({'status': 'success'})
    except Notification.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': 'Notification not found'}, status=404)

@require_POST
@login_required
def dismiss_all_notifications(request):
    try:
        # Update all non-dismissed UserNotifications for the user
        user_notifications = UserNotification.objects.filter(
            tenant=request.user.tenant,
            user=request.user,
            dismissed=False
        )
        updated_count = user_notifications.update(
            dismissed=True,
            seen_at=timezone.now()
        )
        return JsonResponse({'success': True, 'updated_count': updated_count})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
def email_config(request):
    user = CustomUser.objects.get(username=request.user.username, tenant=request.user.tenant)
    if request.method == 'POST':
        email_config_form = EmailConfigForm(request.POST, instance=user)
        if email_config_form.is_valid():
            email_provider = email_config_form.cleaned_data["email_provider"]
            email_address = email_config_form.cleaned_data["email_address"]
            email_password = email_config_form.cleaned_data["email_password"]
            
            # Test SMTP connection
            connection, error_message = get_email_smtp_connection(email_provider, email_address, email_password)
            
            if connection:
                # Save credentials to the user's profile
                user = request.user
                user.email_provider = email_provider
                user.email_address = email_address
                user.email_password = email_password  # Consider encrypting
                user.save()
                logger.info(f"Email config saved for {email_address}")
                return HttpResponseRedirect("/dashboard/email-config/success/")
            else:
                # Handle specific errors (e.g., Gmail 534)
                error_context = {"error": error_message}
                if "534" in error_message and email_provider == "gmail":
                    error_context["provider"] = "Gmail"
                    error_context["help_url"] = "https://support.google.com/mail/?p=InvalidSecondFactor"
                    error_context["message"] = (
                        "Gmail requires an app-specific password for third-party apps. "
                        "Please generate one in your Google Account settings and try again."
                    )
                elif email_provider == "zoho" and "authentication" in error_message.lower():
                    error_context["provider"] = "Zoho"
                    error_context["help_url"] = "https://www.zoho.com/mail/help/app-specific-passwords.html"
                    error_context["message"] = (
                        "Zoho requires an app-specific password for third-party apps. "
                        "Please generate one in your Zoho Account settings and try again."
                    )
                elif email_provider == "outlook" and "535" in error_message:
                    error_context["provider"] = "Outlook"
                    error_context["help_url"] = "https://support.microsoft.com/en-us/office/pop-imap-and-smtp-settings-for-outlook-com-d088b986-291d-42b8-9564-9c414e2aa040"
                    error_context["message"] = (
                        "Outlook requires an app-specific password for third-party apps when two-factor authentication is enabled. "
                        "Please generate one in your Microsoft Account settings and try again."
                    )
                elif email_provider == "yahoo" and "535" in error_message:
                    error_context["provider"] = "Yahoo"
                    error_context["help_url"] = "https://help.yahoo.com/kb/SLN15241.html"
                    error_context["message"] = (
                        "Yahoo requires an app-specific password for third-party apps when two-factor authentication is enabled. "
                        "Please generate one in your Yahoo Account settings and try again."
                    )
                elif email_provider == "icloud" and "535" in error_message:
                    error_context["provider"] = "iCloud"
                    error_context["help_url"] = "https://support.apple.com/en-us/HT204397"
                    error_context["message"] = (
                        "iCloud requires an app-specific password for third-party apps. "
                        "Please generate one in your Apple ID settings and try again."
                    )
                else:
                    error_context["provider"] = email_provider.capitalize()
                    error_context["help_url"] = ""  # Add generic help URL if needed
                    error_context["message"] = "Failed to connect to the email server. Please check your credentials or network."
                
                return render(request, "settings/email_config_error.html", error_context)
                # email_config_form.save()
                # return redirect('view_my_profile')
    else:
        email_config_form = EmailConfigForm(instance=user)
    return render(request, 'settings/email_config.html', {'email_config_form': email_config_form})

def email_config_success_view(request):
    return render(request, "settings/email_config_success.html")

from django.db import models
from rest_framework.authtoken.models import Token
from rest_framework.response import Response

class EventViewSet(viewsets.ModelViewSet):
    serializer_class = EventSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return Event.objects.filter(
            models.Q(created_by=user) | models.Q(participants__user=user),
            tenant=user.tenant
        ).distinct()

    def perform_create(self, serializer):
        print(f"Received data: {self.request.data}")
        event = serializer.save(created_by=self.request.user, tenant=self.request.user.tenant)

    def update(self, request, *args, **kwargs):
        event = self.get_object()
        if event.created_by != request.user:
            return Response({"detail": "You can only edit events you created."}, status=403)
        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        event = self.get_object()
        if event.created_by != request.user:
            return Response({"detail": "You can only delete events you created."}, status=403)
        return super().destroy(request, *args, **kwargs)

from documents.serializers import UserSerializer
class UserViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = CustomUser.objects.all()
    serializer_class = UserSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        # Filter users by tenant
        return CustomUser.objects.filter(tenant=self.request.tenant)
    
from rest_framework.views import APIView
from rest_framework import status

class EventParticipantResponseView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, event_id):
        tenant = request.tenant
        user = request.user
        print("Event User: ", user)
        response = request.data.get('response')

        if response not in ['accepted', 'declined']:
            return Response({'error': 'Invalid response'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            participant = EventParticipant.objects.get(event_id=event_id, user=user, tenant=tenant)
            participant.response = response
            participant.save()
            return Response({'message': 'Response updated successfully'}, status=status.HTTP_200_OK)
        except EventParticipant.DoesNotExist:
            return Response({'error': 'You are not a participant of this event'}, status=status.HTTP_403_FORBIDDEN)


from django.middleware.csrf import get_token

@login_required
def calendar_view(request):
    # auth_token = None
    CustomUser = get_user_model()
    # if request.user.is_authenticated:
    #     auth_token = Token.objects.get(user=request.user).key if Token.objects.filter(user=request.user).exists() else None
    context = {
        # 'auth_token': auth_token or '',
        'csrf_token': get_token(request),
        'notification_bar_items': [],
        'birthday_others': [],
        'birthday_self': False,
        'users': CustomUser.objects.filter(tenant=request.user.tenant),
    }
    print(f"Context: {context}")  # Debug
    return render(request, 'users/calendar.html', context)

@require_POST
def export_staff_csv(request):
    profile_ids = request.POST.getlist('profile_ids')
    if not profile_ids:
        return HttpResponse(
            status=400,
            content_type='application/json',
            content='{"error": "No profiles selected"}'
        )

    # Only export profiles that belong to the same tenant as the requester
    profiles = StaffProfile.objects.filter(
        user__id__in=profile_ids,
        user__tenant=request.user.tenant
    ).select_related('user', 'department').prefetch_related('team')

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="staff_export.csv"'

    writer = csv.writer(response)
    writer.writerow(['Name', 'Phone Number', 'Email', 'Sex', 'Designation', 'Department', 'Team'])

    for profile in profiles:
        writer.writerow([
            f"{profile.first_name} {profile.middle_name or ''} {profile.last_name}".strip(),
            profile.phone_number or 'N/A',
            profile.email or 'N/A',
            profile.sex or 'N/A',
            profile.designation or 'N/A',
            profile.department.name if profile.department else 'N/A',
            ', '.join(profile.team.all().values_list('name', flat=True)) or 'N/A',
        ])

    return response

# ...................................................................................................................

# Admin views

@user_passes_test(is_admin)
def admin_dashboard(request):
    model_links = {
        "Departments": reverse("department_list"),
        "Events": reverse("event_list"),
        "Event Participants": reverse("event_participant_list"),
        "Notifications": reverse("admin_notification_list"),
        "Staff Profiles": reverse("staff_profile_list"),
        "Teams": reverse("admin_team_list"),
        "User Notifications": reverse("user_notification_list"),
        "Users": reverse("users_list"),
    }
    return render(request, "admin/admin_dashboard.html", {"model_links": model_links})

@user_passes_test(is_admin)
def bulk_delete(request, model_name):
    if request.method == "POST":
        # Validate tenant
        if request.user.tenant != request.tenant:
            return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

        # Map model names to actual model classes and their list view names
        model_mapping = {
            'customuser': (CustomUser, 'users_list'),
            'document': (Document, 'admin_document_list'),
            'folder': (Folder, 'admin_folder_list'),
            'file': (File, 'admin_file_list'),
            'task': (Task, 'admin_task_list'),
            'department': (Department, 'department_list'),
            'team': (Team, 'admin_team_list'),
            'event': (Event, 'event_list'),
            'eventparticipant': (EventParticipant, 'event_participant_list'),
            'staffprofile': (StaffProfile, 'staff_profile_list'),
            'notification': (Notification, 'admin_notification_list'),
            'usernotification': (UserNotification, 'user_notification_list'),
        }

        # Check if model_name is valid
        if model_name.lower() not in model_mapping:
            return HttpResponseForbidden("Invalid model name.")

        model_class, redirect_view = model_mapping[model_name.lower()]

        # Get IDs to delete
        ids = request.POST.getlist("ids")
        if not ids:
            return redirect(redirect_view)

        # Delete objects, ensuring they belong to the tenant
        try:
            model_class.objects.filter(id__in=ids, tenant=request.tenant).delete()
        except Exception as e:
            return HttpResponseForbidden(f"Error deleting objects: {str(e)}")

        return redirect(redirect_view)

    return HttpResponseForbidden("Invalid request method.")

@user_passes_test(is_admin)
def bulk_action_users(request):
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")
    
    if request.method != "POST":
        return HttpResponseForbidden("Invalid request method.")
    
    action = request.POST.get('action')
    ids = request.POST.getlist('ids')
    
    if not action or not ids:
        return redirect("users_list")  # Silently redirect if no action or IDs
    
    try:
        users = CustomUser.objects.filter(id__in=ids, tenant=request.tenant)
        if not users.exists():
            return HttpResponseForbidden("No valid users found for this company.")
        
        if action == "delete":
            users.delete()
            LogEntry.objects.log_action(
                user_id=request.user.id,
                content_type_id=ContentType.objects.get_for_model(CustomUser).pk,
                object_id=None,
                object_repr="Multiple users",
                action_flag=CHANGE,
                change_message=f"Bulk deleted {len(ids)} users"
            )
        elif action == "activate":
            updated_count = users.update(is_active=True)
            LogEntry.objects.log_action(
                user_id=request.user.id,
                content_type_id=ContentType.objects.get_for_model(CustomUser).pk,
                object_id=None,
                object_repr="Multiple users",
                action_flag=CHANGE,
                change_message=f"Bulk activated {updated_count} users"
            )
            # Send activation emails
            admin_user = request.user
            sender_provider = admin_user.email_provider
            sender_email = admin_user.email_address
            sender_password = admin_user.email_password
            if sender_email and sender_password:
                connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)
                if settings.DEBUG:
                    base_domain = "localhost:8000"
                    protocol = "http"
                else:
                    base_domain = "teammanager.ng"
                    protocol = "https"
                login_url = f"{protocol}://{request.tenant.slug}.{base_domain}/accounts/login"
                
                for user in users:
                    subject = f"Account Approval: {user.username}"
                    message = f"""
                    Dear {user.username},

                    Your account has been activated. Please click the link below to log in:
                    {login_url}

                    Best regards,  
                    {admin_user.get_full_name() or admin_user.username}
                    """
                    try:
                        send_mail(
                            subject,
                            message,
                            sender_email,
                            [user.email],
                            connection=connection,
                        )
                    except Exception as e:
                        print(f"Failed to send email to {user.email}: {e}")
                        # Continue with other users even if one email fails
        else:
            return HttpResponseForbidden("Invalid action specified.")
    
    except Exception as e:
        return HttpResponseForbidden(f"Error processing action: {str(e)}")
    
    return redirect("users_list")


@user_passes_test(is_admin)
def users_list(request):
    # Validate that the requesting user belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: User does not belong to this company.")

    # Filter users by the current tenant
    users = CustomUser.objects.filter(tenant=request.tenant).order_by('date_joined')
    paginator = Paginator(users, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    count = users.count()
    return render(request, "admin/users_list.html", {"users": page_obj, 'count': count})

@user_passes_test(is_admin)
def create_user(request):
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")
    if request.method == "POST":
        form = UserForm(request.POST, tenant=request.tenant)
        if form.is_valid():
            user = form.save(commit=False)
            user.set_password(form.cleaned_data["password"])
            user.tenant = request.tenant
            user.save()
            LogEntry.objects.log_action(
                user_id=request.user.id,
                content_type_id=ContentType.objects.get_for_model(CustomUser).pk,
                object_id=user.id,
                object_repr=user.username,
                action_flag=ADDITION,
                change_message='Created user'
            )
            return redirect("users_list")
    else:
        form = UserForm(tenant=request.tenant)
    return render(request, "admin/create_user.html", {"form": form})

@user_passes_test(is_admin)
def view_user_details(request, user_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the user, ensuring they belong to the same tenant
    try:
        user_view = CustomUser.objects.get(id=user_id, tenant=request.tenant)
        details = ['username', 'first_name', 'last_name', 'email', 
                 'is_active', 'roles', 'phone_number', 
                  'department', 'teams', 'email_address', 'email_password']
    except CustomUser.DoesNotExist:
        # return HttpResponseForbidden("User not found or does not belong to your tenant.")
        raise PermissionDenied("User not found or does not belong to your company.")

    return render(request, "admin/view_user_details.html", {"user_view": user_view, "details": details})
@user_passes_test(is_admin)
# @user_passes_test(lambda u: u.is_superuser)
def approve_user(request, user_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the user, ensuring they belong to the same tenant
    try:
        user = CustomUser.objects.get(id=user_id, tenant=request.tenant)
    except CustomUser.DoesNotExist:
        return HttpResponseForbidden("User not found or does not belong to your company.")

    # Activate the user
    user.is_active = True
    user.save()

    # Use the admin's email credentials (request.user is already the admin)
    admin_user = request.user
    if admin_user.email_address and admin_user.email_password:
        sender_provider = admin_user.email_provider
        sender_email = admin_user.email_address
        sender_password = admin_user.email_password
    else:
        superuser = CustomUser.objects.get(is_superuser=True)
        sender_provider = superuser.email_provider
        sender_email = superuser.email_address
        sender_password = superuser.email_password

    if not sender_email or not sender_password:
        return HttpResponseForbidden("Your email credentials are missing. Contact admin.")

    # Set up email connection
    connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)

    # Generate tenant-specific login URL
    # Determine base domain based on environment
    if settings.DEBUG:
        base_domain = "127.0.0.1:8000"  # Local development
        protocol = "http"
    else:
        base_domain = "teammanager.ng"  # Production
        protocol = "https"

    # Generate tenant-specific login URL
    login_url = f"{protocol}://{request.tenant.slug}.{base_domain}/accounts/login"

    # Prepare email
    subject = f"Account Approval: {user.username}"
    message = f"""
    Dear {user.username},

    Your account has been activated. Please click the link below to log in:
    {login_url}

    Best regards,  
    {admin_user.get_full_name() or admin_user.username}
    """

    print("Sending mail...")

    try:
        # Send email to the user
        send_mail(
            subject,
            message,
            sender_email,
            [user.email],
            connection=connection,
        )
    except Exception as e:
        print(f"Failed to send email: {e}")
        return HttpResponseForbidden("Failed to send approval email. Contact admin.")

    return redirect("users_list")

@user_passes_test(is_admin)
def edit_user(request, user_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the user, ensuring they belong to the same tenant
    user = get_object_or_404(CustomUser, id=user_id, tenant=request.tenant)

    if request.method == "POST":
        form = EditUserForm(request.POST, instance=user, tenant=request.tenant)
        if form.is_valid():
            # Ensure the tenant field cannot be changed
            form.instance.tenant = request.tenant
            form.save()
            LogEntry.objects.log_action(
                user_id=request.user.id,
                content_type_id=ContentType.objects.get_for_model(CustomUser).pk,
                object_id=user.id,
                object_repr=user.username,
                action_flag=CHANGE,
                change_message='Edited user profile'
            )
            return redirect("users_list")
    else:
        form = EditUserForm(instance=user, tenant=request.tenant)
    return render(request, "admin/edit_user.html", {"form": form})

@user_passes_test(is_admin)
def admin_documents_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    documents = Document.objects.filter(tenant=request.tenant)
    paginator = Paginator(documents, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/documents_list.html", {"documents": page_obj})

@user_passes_test(is_admin)
def admin_document_details(request, document_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the document, ensuring it belongs to the same tenant
    try:
        document_view = Document.objects.get(id=document_id, tenant=request.tenant)
    except Document.DoesNotExist:
        return HttpResponseForbidden("Document not found or does not belong to your company.")

    return render(request, "admin/view_document_details.html", {"document_view": document_view})

@user_passes_test(is_admin)
def admin_delete_document(request, document_id):
    if hasattr(request, 'tenant') and request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized to perform actions for this company.")
    document = get_object_or_404(Document, id=document_id, tenant=request.tenant)

    # Ensure document files are deleted from storage
    if document.word_file:
        document.word_file.delete(save=False)
    if document.pdf_file:
        document.pdf_file.delete(save=False)

    document.delete()
    return redirect("admin_document_list")  # Redirect back to the list

@user_passes_test(is_admin)
def admin_folder_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    folders = Folder.objects.filter(tenant=request.tenant)
    paginator = Paginator(folders, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/folder_list.html", {"folders": page_obj})

@user_passes_test(is_admin)
def admin_folder_details(request, folder_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the folder, ensuring it belongs to the same tenant
    try:
        folder_view = Folder.objects.get(id=folder_id, tenant=request.tenant)
    except Folder.DoesNotExist:
        return HttpResponseForbidden("Folder not found or does not belong to your company.")

    return render(request, "admin/view_folder_details.html", {"folder_view": folder_view})

@user_passes_test(is_admin)
def admin_delete_folder(request, folder_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the folder, ensuring it belongs to the same tenant
    folder = get_object_or_404(Folder, id=folder_id, created_by=request.user, tenant=request.tenant)
    folder.delete()
    return redirect("admin_folder_list")

@user_passes_test(is_admin)
def admin_file_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    files = File.objects.filter(tenant=request.tenant)
    paginator = Paginator(files, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/file_list.html", {"files": page_obj})

@user_passes_test(is_admin)
def admin_delete_file(request, file_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the file, ensuring it belongs to the same tenant
    file = get_object_or_404(File, id=file_id, uploaded_by=request.user, tenant=request.tenant)
    file.delete()
    return redirect("admin_file_list")

@user_passes_test(is_admin)
def admin_task_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    tasks = Task.objects.filter(tenant=request.tenant)
    paginator = Paginator(tasks, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/task_list.html", {"tasks": page_obj})

@user_passes_test(is_admin)
def admin_task_detail(request, task_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the task, ensuring it belongs to the same tenant
    try:
        task_view = Task.objects.get(id=task_id, tenant=request.tenant)
    except Task.DoesNotExist:
        return HttpResponseForbidden("Task not found or does not belong to your company.")

    return render(request, "admin/view_task_details.html", {"task_view": task_view})

@user_passes_test(is_admin)
def department_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")
    
    # Fetch all departments in that tenant
    departments = Department.objects.filter(tenant=request.tenant)
    paginator = Paginator(departments, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/department_list.html", {"departments": page_obj})

@user_passes_test(is_admin)
def create_department(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    if request.method == "POST":
        form = DepartmentForm(request.POST, user=request.user)
        if form.is_valid():
            department = form.save(commit=False)
            department.tenant = request.tenant
            department.save()
            department.hod.department = department
            department.hod.save()
            return redirect("department_list")
    else:
        form = DepartmentForm(user=request.user)
    return render(request, "admin/create_department.html", {"form": form})

@user_passes_test(is_admin)
def edit_department(request, department_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the department, ensuring it belongs to the same tenant
    department = get_object_or_404(Department, id=department_id, tenant=request.tenant)

    if request.method == "POST":
        form = DepartmentForm(request.POST, instance=department, user=request.user)
        if form.is_valid():
            department=form.save(commit=False)
            department.hod.department = department
            department.hod.save()
            department.save()
            return redirect("department_list")
    else:
        form = DepartmentForm(instance=department, user=request.user)
    return render(request, "admin/edit_department.html", {"form": form})

@user_passes_test(is_admin)
def delete_department(request, department_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the department, ensuring it belongs to the same tenant
    department = get_object_or_404(Department, id=department_id, tenant=request.tenant)
    department.delete()
    return redirect("department_list")

@user_passes_test(is_admin)
def admin_team_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    teams = Team.objects.filter(tenant=request.tenant).order_by('department')
    paginator = Paginator(teams, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/team_list.html", {"teams": page_obj})

@user_passes_test(is_admin)
def create_team(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    if request.method == "POST":
        form = TeamForm(request.POST, user=request.user)
        if form.is_valid():
            team = form.save(commit=False)
            team.tenant = request.tenant
            team.save()
            return redirect("admin_team_list")
    else:
        form = TeamForm(user=request.user)
    return render(request, "admin/create_team.html", {"form": form})

@user_passes_test(is_admin)
def delete_team(request, team_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the team, ensuring it belongs to the same tenant
    team = get_object_or_404(Team, id=team_id, tenant=request.tenant)
    team.delete()
    return redirect("admin_team_list")

@user_passes_test(is_admin)
def edit_team(request, team_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")
    
    # Get the team, ensuring it belongs to the same tenant
    team = get_object_or_404(Team, id=team_id, tenant=request.tenant)

    if request.method == "POST":
        form = TeamForm(request.POST, instance=team, user=request.user)
        if form.is_valid():
            form.save()
            return redirect("admin_team_list")
    else:
        form = TeamForm(instance=team, user=request.user)
    return render(request, "admin/edit_team.html", {"form": form})

@user_passes_test(is_admin)
def event_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    events = Event.objects.filter(tenant=request.tenant)
    paginator = Paginator(events, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/event_list.html", {"events": page_obj})

@user_passes_test(is_admin)
def create_event(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    if request.method == "POST":
        form = EventForm(request.POST)
        if form.is_valid():
            event = form.save(commit=False)
            event.tenant = request.tenant
            event.created_by = request.user
            event.save()
            return redirect("event_list")
    else:
        form = EventForm()
    return render(request, "admin/create_event.html", {"form": form})

@user_passes_test(is_admin)
def edit_event(request, event_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    event = get_object_or_404(Event, id=event_id, tenant=request.tenant)

    if request.method == "POST":
        form = EventForm(request.POST, instance=event, user=request.user)
        if form.is_valid():
            form.save()
            return redirect("event_list")
    else:
        form = EventForm(instance=event, user=request.user)
    return render(request, "admin/edit_event.html", {"form": form})

@user_passes_test(is_admin)
def delete_event(request, event_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    event = get_object_or_404(Event, id=event_id, tenant=request.tenant)
    event.delete()
    return redirect("event_list")

@user_passes_test(is_admin)
def event_participant_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    participants = EventParticipant.objects.filter(tenant=request.tenant).order_by('event')
    paginator = Paginator(participants, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/event_participant_list.html", {"participants": page_obj})

@user_passes_test(is_admin)
def create_event_participant(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    if request.method == "POST":
        form = EventParticipantForm(request.POST, user=request.user)
        if form.is_valid():
            event_participant=form.save(commit=False)
            event_participant.tenant = request.tenant
            event_participant.save()
            return redirect("event_participant_list")
    else:
        form = EventParticipantForm(user=request.user)
    return render(request, "admin/create_event_participant.html", {"form": form})

@user_passes_test(is_admin)
def edit_event_participant(request, event_participant_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    event_participant = get_object_or_404(EventParticipant, id=event_participant_id, tenant=request.tenant)

    if request.method == "POST":
        form = EventParticipantForm(request.POST, instance=event_participant, user=request.user)
        if form.is_valid():
            form.save()
            return redirect("event_participant_list")
    else:
        form = EventParticipantForm(instance=event_participant, user=request.user)
    return render(request, "admin/edit_event_participant.html", {"form": form})

@user_passes_test(is_admin)
def delete_event_participant(request, event_participant_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    event_participant = get_object_or_404(EventParticipant, id=event_participant_id, tenant=request.tenant)
    event_participant.delete()
    return redirect("event_participant_list")

@user_passes_test(is_admin)
def staff_profile_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    staff_profiles = StaffProfile.objects.filter(tenant=request.tenant)
    paginator = Paginator(staff_profiles, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/staff_profile_list.html", {"staff_profiles": page_obj})

@user_passes_test(is_admin)
def create_staff_profile(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    if request.method == "POST":
        form = StaffProfileForm(request.POST, user=request.user)
        if form.is_valid():
            staff_profile=form.save(commit=False)
            staff_profile.tenant = request.tenant
            staff_profile.save()
            return redirect("staff_profile_list")
    else:
        form = StaffProfileForm(user=request.user)
    return render(request, "admin/create_staff_profile.html", {"form": form})

@user_passes_test(is_admin)
def edit_staff_profile(request, staff_profile_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    staff_profile = get_object_or_404(StaffProfile, id=staff_profile_id, tenant=request.tenant)

    if request.method == "POST":
        form = StaffProfileForm(request.POST, instance=staff_profile, user=request.user)
        if form.is_valid():
            form.save()
            return redirect("staff_profile_list")
    else:
        form = StaffProfileForm(instance=staff_profile, user=request.user)
    return render(request, "admin/edit_staff_profile.html", {"form": form})

@user_passes_test(is_admin)
def delete_staff_profile(request, staff_profile_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    staff_profile = get_object_or_404(StaffProfile, id=staff_profile_id, tenant=request.tenant)
    staff_profile.delete()
    return redirect("staff_profile_list")

@user_passes_test(is_admin)
def admin_notification_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    notifications = Notification.objects.filter(tenant=request.tenant).order_by('created_at')
    paginator = Paginator(notifications, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/admin_notification_list.html", {"notifications": page_obj})

@user_passes_test(is_admin)
def create_notification(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    if request.method == "POST":
        form = NotificationForm(request.POST)
        if form.is_valid():
            notification = form.save(commit=False)
            notification.tenant = request.tenant
            notification.save()
            return redirect("admin_notification_list")
    else:
        form = NotificationForm()
    return render(request, "admin/create_notification.html", {"form": form})

@user_passes_test(is_admin)
def edit_notification(request, notification_id):
    # Validate taht the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")
    
    notification = get_object_or_404(Notification, id=notification_id, tenant=request.tenant)

    if request.method == "POST":
        form = NotificationForm(request.POST, instance=notification)
        if form.is_valid():
            form.save()
            return redirect("admin_notification_list")
    else:
        form = NotificationForm(instance=notification)
    return render(request, "admin/edit_notification.html", {"form": form})

@user_passes_test(is_admin)
def delete_notification(request, notification_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    notification = get_object_or_404(Notification, id=notification_id, tenant=request.tenant)
    notification.delete()
    return redirect("admin_notification_list")

@user_passes_test(is_admin)
def user_notification_list(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    user_notifications = UserNotification.objects.filter(tenant=request.tenant)
    paginator = Paginator(user_notifications, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, "admin/user_notification_list.html", {"user_notifications": page_obj})

@user_passes_test(is_admin)
def create_user_notification(request):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    if request.method == "POST":
        form = UserNotificationForm(request.POST, user=request.user)
        if form.is_valid():
            user_notification = form.save(commit=False)
            user_notification.tenant = request.tenant
            user_notification.save()
            return redirect("admin_notification_list")
    else:
        form = UserNotificationForm(user=request.user)
    return render(request, "admin/create_user_notification.html", {"form": form})

@user_passes_test(is_admin)
def edit_user_notification(request, user_notification_id):
    # Validate taht the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")
    
    user_notification = get_object_or_404(UserNotification, id=user_notification_id, tenant=request.tenant)

    if request.method == "POST":
        form = UserNotificationForm(request.POST, instance=user_notification, user=request.user)
        if form.is_valid():
            form.save()
            return redirect("user_notification_list")
    else:
        form = UserNotificationForm(instance=user_notification, user=request.user)
    return render(request, "admin/edit_user_notification.html", {"form": form})

@user_passes_test(is_admin)
def delete_user_notification(request, user_notification_id):
    # Validate that the admin belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: Admin does not belong to this company.")

    # Get the event, ensuring it belongs to the same tenant
    user_notification = get_object_or_404(UserNotification, id=user_notification_id, tenant=request.tenant)
    user_notification.delete()
    return redirect("user_notification_list")


@login_required
def view_company_profile(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    print(f"User tenant: {request.user.tenant}. Request tenant: {request.tenant.name}")
    try: 
        tenant_profile, created = CompanyProfile.objects.get_or_create(
            tenant=request.tenant
        )

        num_staff = CustomUser.objects.filter(tenant=request.tenant).count()
        num_departments = Department.objects.filter(tenant=request.tenant).count()
        num_teams = Team.objects.filter(tenant=request.tenant).count()

        tenant_profile.num_staff = num_staff
        tenant_profile.num_departments = num_departments
        tenant_profile.num_teams = num_teams
        tenant_profile.save()

        depts = Department.objects.filter(tenant=request.tenant)
        teams = Team.objects.filter(tenant=request.tenant)
        return render(request, 'admin/company_profile.html', {'tenant_profile': tenant_profile, 'depts': depts, 'teams': teams})
    except Exception as e:
        logger.error(f"Error in view_company_profile: {e}")
        # return render(request, 'tenant_error.html', {'message': 'An unexpected error occurred'}, status=500)
        return HttpResponse("An unexpected error occurred")

@login_required
@user_passes_test(is_admin)
def edit_company_profile(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    company_profile, created = CompanyProfile.objects.get_or_create(
        tenant=request.tenant,
        defaults={'company_name': request.tenant.name}
    )
    
    if request.method == "POST":
        form = CompanyProfileForm(request.POST, request.FILES, instance=company_profile)
        if form.is_valid():
            form.save()
            return redirect("view_company_profile")
    else:
        form = CompanyProfileForm(instance=company_profile)
        document_form = CompanyDocumentForm()
    return render(request, "admin/edit_company_profile.html", {"form": form, 'profile': company_profile, 'document_form': document_form})

@login_required
@user_passes_test(is_admin)
def add_company_document(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    if request.method == 'POST':
        form = CompanyDocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.tenant = request.user.tenant
            document.company_profile = get_object_or_404(CompanyProfile, tenant=request.tenant)
            document.save()
            return JsonResponse({
                'success': True,
                'document': {
                    'id': document.id,
                    'description': document.description or document.document_type,
                    'file_url': document.file.url,
                    'document_type': document.get_document_type_display(),
                    'uploaded_at': document.uploaded_at.strftime('%B %d, %Y')
                }
            })
        else:
            return JsonResponse({
                'success': False,
                'errors': form.errors
            }, status=400)
    return JsonResponse({'success': False, 'error': 'Invalid request'}, status=400)

@login_required
@user_passes_test(is_admin)
def delete_company_document(request, document_id):
    if hasattr(request, 'tenant') and request.user.tenant != request.tenant:
        logger.error(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    
    try:
        # Get the user's StaffProfile (assuming one profile per user)
        company_profile = CompanyProfile.objects.get(tenant=request.tenant)
        document = get_object_or_404(CompanyDocument, id=document_id, company_profile=company_profile)
        if request.method == 'POST':
            document.delete()
            return JsonResponse({'success': True})
        return JsonResponse({'success': False, 'error': 'Invalid request method'}, status=400)
        # raise BadRequest('Invalid request method')
    except StaffProfile.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Staff profile not found'}, status=404)
    except StaffDocument.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Document not found or not owned by user'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)



# Contact List, Billing and Mass Mailing Section

@login_required
def contact_list(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    contact_list_dept = Contact.objects.filter(tenant=request.tenant, department=request.user.department, is_public=True)
    contact_list_personal = Contact.objects.filter(tenant=request.tenant, created_by=request.user)
    contact_list = contact_list_dept | contact_list_personal
    paginator = Paginator(contact_list, 10)  # 10 users per page
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    paginator_dept = Paginator(contact_list_dept, 10)  # 10 users per page
    page_obj_dept = paginator_dept.get_page(page_number)
    paginator_personal = Paginator(contact_list_personal, 10)  # 10 users per page
    page_obj_personal = paginator_personal.get_page(page_number)
    return render(request, 'dashboard/contact_list.html', {'contact_list': page_obj, 'contact_list_dept': page_obj_dept, 'contact_list_personal': page_obj_personal})

@login_required
def create_contact(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        # return HttpResponseForbidden("You are not authorized for this tenant.")
        return render(request, 'error.html', {'message': 'You are not authorized for this company.'})
    if request.method == "POST":
        form = ContactForm(request.POST)
        if form.is_valid():
            contact = form.save(commit=False)
            contact.tenant = request.tenant
            contact.department = request.user.department
            contact.team = request.user.teams.first()
            contact.created_by = request.user
            contact.save()
            return redirect("contact_list")
    else:
        form = ContactForm()
    return render(request, "dashboard/create_contact.html", {"form": form})

@login_required
def view_contact_detail(request, contact_id):
    contact = get_object_or_404(Contact, id=contact_id, tenant=request.tenant)
    data = {
        'id': contact.id,
        'name': contact.name,
        'email': contact.email,
        'phone': contact.phone or '',
        'organization': contact.organization or '',
        'designation': contact.designation or '',
        'priority': contact.priority or '',
        'department': contact.department.name if contact.department else None,
        'team': contact.team.name if contact.team else None,
        'is_public': contact.is_public,
        'created_by': contact.created_by.username,
        'created_at': contact.created_at.isoformat(),
        'updated_by': contact.updated_by.username if contact.updated_by else None,
        'updated_at': contact.updated_at.isoformat()
    }
    return JsonResponse(data)

@login_required
def edit_contact(request, contact_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        # return HttpResponseForbidden("You are not authorized for this tenant.")
        return render(request, 'error.html', {'message': 'You are not authorized for this company.'})
    contact = get_object_or_404(Contact, id=contact_id, tenant=request.tenant)
    print(f"Contact to edit: {contact}")
    if request.method == "POST":
        form = ContactForm(request.POST, instance=contact)
        if form.is_valid():
            contact = form.save(commit=False)
            contact.tenant = request.tenant
            contact.updated_by = request.user
            contact.save()
            return redirect("contact_list")
    else:
        form = ContactForm(instance=contact)
    return render(request, "dashboard/edit_contact.html", {"form": form})

@login_required
def delete_contact(request, contact_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        # return HttpResponseForbidden("You are not authorized for this tenant.")
        return render(request, 'error.html', {'message': 'You are not authorized for this company.'})
    contact = get_object_or_404(Contact, id=contact_id, tenant=request.tenant)
    if contact.created_by != request.user:
        return JsonResponse({'success': False, 'errors': {'folder': ['You can only delete your own contacts']}}, status=403)
    contact.delete()
    return redirect("contact_list")

@login_required
def email_list(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        # return HttpResponseForbidden("You are not authorized for this tenant.")
        return render(request, 'error.html', {'message': 'You are not authorized for this company.'})
        
    email_list = Email.objects.filter(tenant=request.tenant, sender=request.user)
    email_list_draft = Email.objects.filter(tenant=request.tenant, sender=request.user, sent=False)
    email_list_sent = Email.objects.filter(tenant=request.tenant, sender=request.user, sent=True)
    page_number = request.GET.get('page')
    paginator = Paginator(email_list, 10)  # 10 emails per page
    page_obj = paginator.get_page(page_number)
    paginator_draft = Paginator(email_list_draft, 10)  # 10 emails per page
    page_obj_draft = paginator_draft.get_page(page_number)
    paginator_sent = Paginator(email_list_sent, 10)  # 10 emails per page
    page_obj_sent = paginator_sent.get_page(page_number)
    return render(request, 'dashboard/email_list.html', {'email_list': page_obj, 'email_list_draft': page_obj_draft, 'email_list_sent': page_obj_sent})

@login_required
def edit_email(request, email_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'error.html', {'message': 'You are not authorized for this company.'})

    email = get_object_or_404(Email, id=email_id, tenant=request.tenant, sender=request.user)
    if email.sender != request.user:
        return render(request, 'error.html', {'message': 'You can only edit your own emails.'})
    if email.sent:
        return redirect('email_detail', email_id=email_id)
    
    attachments = email.attachments.all()

    if request.method == 'POST':
        form = EmailForm(request.POST, request.FILES, instance=email)
        
        if form.is_valid():
            email = form.save(commit=False)
            email.tenant = request.tenant
            email.sender = request.user
            email.sent = False
            email.save()

            # Handle attachment deletions if specified
            if 'delete_attachments' in request.POST:
                attachment_ids = request.POST.getlist('delete_attachments')
                Attachment.objects.filter(id__in=attachment_ids, email=email).delete()

            # Save new attachments (handled by form.save())
            form.save()

            # If user wants to send the email
            if 'send' in request.POST:
                sender_provider = request.user.email_provider
                sender_email = request.user.email_address
                sender_password = request.user.email_password
                connection, error_message = get_email_smtp_connection(sender_provider,sender_email, sender_password)
                email_msg = EmailMessage(
                    subject=email.subject,
                    body=email.body,
                    from_email=sender_email,
                    to=email.get_to_emails(),
                    cc=email.get_cc_emails(),
                    bcc=email.get_bcc_emails(),
                    connection=connection
                )
                # Attach all files
                for attachment in email.attachments.all():
                    email_msg.attach_file(attachment.file.path)
                email_msg.send()
                email.sent = True
                email.sent_at = timezone.now()
                email.save()
                return redirect('email_list')

            return redirect('email_list')
    else:
        form = EmailForm(instance=email)

    return render(request, 'dashboard/edit_email.html', {'form': form, 'email': email, 'attachments': attachments})

@login_required
def save_draft(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'error.html', {'message': 'You are not authorized for this tenant.'})

    if request.method == 'POST':
        form = EmailForm(request.POST, request.FILES)
        
        if form.is_valid():
            email = form.save(commit=False)
            email.tenant = request.tenant
            email.sender = request.user
            email.sent = False
            email.save()
            form.save()  # Save attachments
            return redirect('email_list')
    else:
        form = EmailForm()

    return render(request, 'dashboard/send_email.html', {'form': form})

# Update send_email view to handle multiple attachments
@login_required
def send_email(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'error.html', {'message': 'You are not authorized for this tenant.'})

    sender_provider = request.user.email_provider
    sender_email = request.user.email_address
    sender_password = request.user.email_password
    connection, error_message = get_email_smtp_connection(sender_provider,sender_email, sender_password)

    if request.method == 'POST':
        form = EmailForm(request.POST, request.FILES)
        
        if form.is_valid():
            email = form.save(commit=False)
            email.tenant = request.tenant
            email.sender = request.user
            email.sent = False
            email.save()
            form.save()  # Save attachments

            # Send email
            email_msg = EmailMessage(
                subject=email.subject,
                body=email.body,
                from_email=sender_email,
                to=email.get_to_emails(),
                cc=email.get_cc_emails(),
                bcc=email.get_bcc_emails(),
                connection=connection
            )
            for attachment in email.attachments.all():
                email_msg.attach_file(attachment.file.path)
            email_msg.send()
            email.sent = True
            email.sent_at = timezone.now()
            email.save()
            return redirect('email_list')
    else:
        form = EmailForm()

    return render(request, 'dashboard/send_email.html', {'form': form})
@login_required
def email_detail(request, email_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    email = get_object_or_404(Email, id=email_id, tenant=request.tenant, sender=request.user)
    return render(request, 'dashboard/email_detail.html', {'email': email})

def delete_email(request, email_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return HttpResponseForbidden("You are not authorized for this company.")
    email = get_object_or_404(Email, id=email_id, tenant=request.tenant, sender=request.user)
    if email.sender != request.user:
        raise HttpResponseForbidden('You can only delete your own emails')
    email.delete()
    return redirect('email_list')

@login_required
def delete_email_attachment(request, email_id, attachment_id):
    email = get_object_or_404(Email, id=email_id, tenant=request.tenant)
    attachment = get_object_or_404(Attachment, id=attachment_id, tenant=request.tenant)
    if not (email.sender == request.user) or attachment not in email.attachmets.all():
        return JsonResponse({'error': 'Access denied'}, status=403)
    
    if request.method == 'POST':
        email.attachments.remove(attachment)
        if not email.attachments.exists():
            attachment.delete()
        return JsonResponse({'success': True})
    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required
def contact_search(request):
    query = request.GET.get('q', '')
    contacts = Contact.objects.filter(
        tenant=request.user.tenant,  # Assuming tenant-based filtering
        email__icontains=query
    )[:10]  # Limit to 10 results
    results = [{'email': contact.email, 'name': contact.name} for contact in contacts]
    return JsonResponse(results, safe=False)

def contact_support(request):
    if request.method == 'POST':
        print("POST data: %s", dict(request.POST))
        files = request.FILES.getlist('attachments')
        print("FILES data: %s", request.FILES)
        print(("FILES data: %s", [(f.name, f.size, f.content_type) for f in files] if files else "No files received"))
        form = SupportForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                # Create email
                email = EmailMessage(
                    subject=form.cleaned_data['subject'],
                    body=form.cleaned_data['message'],
                    from_email=request.user.email_address if request.user.is_authenticated else 'no-reply@teammanager.ng',
                    to=['faith.osebi@transnetcloud.com'],
                    connection=get_email_smtp_connection(request.user.email_provider, request.user.email_address, request.user.email_password)
                )
                
                # Handle attachments
                for f in request.FILES.getlist('attachments'):
                    print(("Attaching file: %s (size: %s, type: %s)", f.name, f.size, f.content_type))
                    email.attach(f.name, f.read(), f.content_type)

                # for f in form.cleaned_data.get('attachments', []):
                #     print(("Attaching file: %s (size: %s, type: %s)", f.name, f.size, f.content_type))
                #     email.attach(f.name, f.read(), f.content_type)
                
                # Send email
                email.send()
                return JsonResponse({'success': True})
            except Exception as e:
                print(("Email sending error: %s", str(e)))
                return JsonResponse({'success': False, 'error': str(e)})
        else:
            print(("Form errors: %s", form.errors))
            return JsonResponse({'success': False, 'error': 'Invalid form data'})
    else:
        form = SupportForm()
    return render(request, 'dashboard/contact_support.html', {'form': form})

# HR
# VACANCIES

@login_required
@user_passes_test(is_hr)
def hr_dashboard(request):
    return render(request, 'hr/hr_dashboard.html')

@login_required
@user_passes_test(is_hr)
def vacancy_list(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancies = Vacancy.objects.filter(tenant=request.tenant).select_related('created_by', 'updated_by', 'shared_by')
    for vacancy in vacancies:
        if vacancy.share_time_end:
            if vacancy.is_shared and vacancy.share_time <= timezone.now() <= vacancy.share_time_end:
                vacancy.shareable_link = request.build_absolute_uri(vacancy.get_shareable_link())
        if vacancy.is_shared and vacancy.share_time <= timezone.now():
            vacancy.shareable_link = request.build_absolute_uri(vacancy.get_shareable_link()) 
        else:
            vacancy.shareable_link = None

    paginator = Paginator(vacancies, 10)  # 10 vacancies per page
    page = request.GET.get('page')
    page_obj = paginator.get_page(page)
    
    return render(request, 'hr/vacancy_list.html', {'vacancies': page_obj})

@login_required
@user_passes_test(is_hr)
def create_vacancy(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    if request.method == 'POST':
        form = VacancyForm(request.POST)
        if form.is_valid():
            vacancy = form.save(commit=False)
            vacancy.tenant = request.tenant
            vacancy.created_by = request.user
            vacancy.save()
            return redirect('vacancy_list')
    else:
        form = VacancyForm()
    return render(request, 'hr/create_vacancy.html', {'form': form}) 

@login_required
@user_passes_test(is_hr)
def edit_vacancy(request, vacancy_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id, tenant=request.tenant)
    if request.method == 'POST':
        form = VacancyForm(request.POST, instance=vacancy)
        if form.is_valid():
            form.save(commit=False)
            form.updated_by = request.user
            form.updated_at = timezone.now()
            form.save()
            return redirect('vacancy_list')
    else:
        form = VacancyForm(instance=vacancy)
    return render(request, 'hr/edit_vacancy.html', {'form': form})

@login_required
@user_passes_test(is_hr)
def vacancy_detail(request, vacancy_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id, tenant=request.tenant)
    
    return render(request, 'hr/vacancy_detail.html', {'vacancy': vacancy})

@login_required
@user_passes_test(is_hr)
def delete_vacancy(request, vacancy_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id, tenant=request.tenant)
    vacancy.delete()
    return redirect('vacancy_list')

@login_required
@user_passes_test(is_hr)
def share_vacancy(request, vacancy_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id, tenant=request.tenant)
    if request.method == 'POST':
        end_date = request.POST.get('end_date')
        try:
            share_time_end = timezone.datetime.strptime(end_date, '%Y-%m-%d') if end_date else None
            if share_time_end:
                share_time_end = timezone.make_aware(share_time_end)
        except ValueError:
            return JsonResponse({"success": False, "error": "Invalid date format. Use YYYY-MM-DD."}, status=400)

        vacancy.is_shared = True
        vacancy.status = 'active'
        vacancy.shared_by = request.user
        vacancy.share_time = timezone.now()
        vacancy.share_time_end = share_time_end
        vacancy.save()
        vacancy.shareable_link = request.build_absolute_uri(vacancy.get_shareable_link())
        return JsonResponse({"success": True, "vacancy": vacancy.id, "shareable_link": vacancy.shareable_link})
    
    return JsonResponse({"success": False, "error": "Method not allowed"}, status=405)

@login_required
@user_passes_test(is_hr)
def withdraw_vacancy(request, vacancy_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    vacancy = get_object_or_404(Vacancy, id=vacancy_id, tenant=request.tenant)
    vacancy.is_shared = False
    vacancy.shared_by = None
    vacancy.share_time = None
    vacancy.share_time_end = None
    vacancy.save()
    return JsonResponse({"success": True})

def vacancy_post(request, token):
    vacancy = get_object_or_404(Vacancy, share_token=token)
    if not vacancy.is_shared or (vacancy.share_time_end and timezone.now() > vacancy.share_time_end) or vacancy.status in ['closed', 'withdrawn']:
        message = "This vacancy is no longer available."
        if vacancy.status == 'closed':
            message = "This vacancy is closed."
        elif vacancy.status == 'withdrawn':
            message = "This vacancy is withdrawn."
        return render(request, 'hr/vacancy_expired.html', {'message': message})
    
    return render(request, 'hr/vacancy_post.html', {'vacancy': vacancy})

# Vacancy Application

@login_required
@user_passes_test(is_hr)
def vacancy_application_list(request):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = Vacancy.objects.filter(tenant=request.tenant)
    paginator = Paginator(vacancy, 10)  # 10 applications per page
    page = request.GET.get('page')
    page_obj = paginator.get_page(page)
    
    return render(request, 'hr/vacancy_application_list.html', {'vacancies': page_obj})


def create_vacancy_application(request, vacancy_id):
    vacancy = get_object_or_404(Vacancy, id=vacancy_id)
    if request.method == 'POST':
        form = VacancyApplicationForm(request.POST, request.FILES)
        if form.is_valid():
            vacancy_application = form.save(commit=False)
            vacancy_application.vacancy = vacancy
            vacancy_application.tenant = request.tenant
            vacancy_application.save()
            return render(request, 'hr/vacancy_application_success.html', {'name':vacancy_application.first_name, 'vacancy': vacancy})
    else:
        form = VacancyApplicationForm()
    return render(request, 'hr/create_vacancy_application.html', {'form': form, 'vacancy': vacancy}) 

@login_required
@user_passes_test(is_hr)
def applications_per_vacancy(request, vacancy_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id, tenant=request.tenant)
    vacancy_applications = VacancyApplication.objects.filter(tenant=request.tenant, vacancy=vacancy)
    paginator = Paginator(vacancy_applications, 10)  # 10 applications per page
    page = request.GET.get('page')
    page_obj = paginator.get_page(page)
    
    return render(request, 'hr/applications_per_vacancy.html', {'vacancy_applications': page_obj, 'vacancy': vacancy})

@login_required
@user_passes_test(is_hr)
def vacancy_application_detail(request, vacancy_id, application_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id)
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant, vacancy=vacancy)
    
    return render(request, 'hr/vacancy_application_detail.html', {'vacancy_application': vacancy_application})

@login_required
@user_passes_test(is_hr)
def delete_vacancy_application(request, vacancy_id, application_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id)
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant, vacancy=vacancy)
    vacancy_application.delete()
    return redirect('applications_per_vacancy', vacancy_id=vacancy_id)

def send_vacancy_accepted_mail(request, application_id):
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant)
    vacancy = vacancy_application.vacancy
    hr = request.user
    if not is_hr(hr):
        print(f"Unauthorized access by user {request.user.username}. Only HRs can perform this action")
        return render(request, 'tenant_error.html', {'error_code': '403','message': 'Only HRs can perform this action.'})
    sender_provider = hr.email_provider if hr.email_provider else SUPERUSER_EMAIL_PROVIDER
    sender_email = hr.email_address if hr.email_address else SUPERUSER_EMAIL_ADDRESS
    sender_password = hr.email_password if hr.email_password else SUPERUSER_EMAIL_PASSWORD
    candidate_name = vacancy_application.first_name
    company = vacancy_application.tenant


    if not sender_email or not sender_password:
        return HttpResponseForbidden("Your email credentials are missing. Contact admin.")

    connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)

    subject = f"You're Moving Forward! Next Steps for the {vacancy.title} Role"
    message = f"""
    Dear {candidate_name},

    Congratulations! Thank you for your application for the {vacancy.title} position at {company}. We were thoroughly impressed with your background, and we are excited to inform you that you have been selected to move forward in our hiring process!

    Your next step is to stay tuned. Our team is currently coordinating the next phase, and you will be receiving a follow-up email from us shortly with detailed instructions.

    Please be sure to keep a close eye on your email inbox (including your spam or promotions folder) so you don't miss our message.

    We are very much looking forward to connecting with you soon and learning more about how you could contribute to our team at {company}.

    If you have any immediate questions, please feel free to reply to this email.

    Best regards,

    Human Resouces Team, 
    {company}
    """

    print("Sending mail...")

    # Create email with attachment
    email = EmailMessage(subject, message, sender_email, [vacancy_application.email], connection=connection)
    email.send()
    
    print("Mail Sent")


def send_vacancy_rejected_mail(request, application_id):
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant)
    vacancy = vacancy_application.vacancy
    hr = request.user
    if not is_hr(hr):
        print(f"Unauthorized access by user {request.user.username}. Only HRs can perform this action")
        return render(request, 'tenant_error.html', {'error_code': '403','message': 'Only HRs can perform this action.'})
    
    sender_provider = hr.email_provider if hr.email_provider else SUPERUSER_EMAIL_PROVIDER
    sender_email = hr.email_address if hr.email_address else SUPERUSER_EMAIL_ADDRESS
    sender_password = hr.email_password if hr.email_password else SUPERUSER_EMAIL_PASSWORD
    candidate_name = vacancy_application.first_name.capitalize()
    company = vacancy_application.tenant.name


    if not sender_email or not sender_password:
        return HttpResponseForbidden("Your email credentials are missing. Contact admin.")

    connection, error_message = get_email_smtp_connection(sender_provider, sender_email, sender_password)

    subject = f"An Update on Your Application for {vacancy.title} Role"
    message = f"""
    Dear {candidate_name},

    Thank you for taking the time to apply for the {vacancy.title} position at {company} and for your interest in joining our team.

    We appreciate the opportunity to learn about your skills and accomplishments. After careful review, we have decided to move forward with candidates whose experience more closely aligns with the specific requirements of this role.

    This was a difficult decision due to the high volume of qualified applicants we received.

    We encourage you to keep an eye on our careers page for future opportunities that may be a better fit for your background

    We wish you the very best in your job search and future endeavors.

    Sincerely,

    Human Resouces Team, 
    {company}
    """

    print("Sending mail...")

    # Create email with attachment
    email = EmailMessage(subject, message, sender_email, [vacancy_application.email], connection=connection)
    email.send()
    
    print("Mail Sent")

@login_required
@user_passes_test(is_hr)
def accept_vacancy_application(request, vacancy_id, application_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id)
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant, vacancy=vacancy)
    if request.method == 'POST':
        status = request.POST.get('accepted')
        if status == 'accepted':
            vacancy_application.status = status
            vacancy_application.save()
            send_vacancy_accepted_mail(request, application_id)
            return JsonResponse({"success": True, "vacancy_id": vacancy_application.id})
    
    return JsonResponse({"success": False, "error": "Method not allowed"}, status=405)

@login_required
@user_passes_test(is_hr)
def reject_vacancy_application(request, vacancy_id, application_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy = get_object_or_404(Vacancy, id=vacancy_id)
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant, vacancy=vacancy)
    if request.method == 'POST':
        status = request.POST.get('rejected')
        if status == 'rejected':
            vacancy_application.status = status
            vacancy_application.save()
            send_vacancy_rejected_mail(request, application_id)
            return JsonResponse({"success": True, "vacancy_id": vacancy_application.id})
    
    return JsonResponse({"success": False, "error": "Method not allowed"}, status=405)
    
# non-form accept, reject vacancy application
@login_required
@user_passes_test(is_hr)
def accept_vac_app(request, application_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant)
    vacancy_application.status = 'accepted'
    vacancy_application.save()
    send_vacancy_accepted_mail(request, application_id)
    return redirect('vacancy_application_detail', vacancy_id=vacancy_application.vacancy.id, application_id=vacancy_application.id)

@login_required
@user_passes_test(is_hr)
def reject_vac_app(request, application_id):
    if not hasattr(request, 'tenant') or request.user.tenant != request.tenant:
        print(f"Unauthorized access by user {request.user.username}: tenant mismatch")
        return render(request, 'tenant_error.html', {'error_code': '401','message': 'You are not authorized for this company.'})
    
    vacancy_application = get_object_or_404(VacancyApplication, id=application_id, tenant=request.tenant)
    vacancy_application.status = 'rejected'
    vacancy_application.save()
    send_vacancy_rejected_mail(request, application_id)
    return redirect('vacancy_application_detail', vacancy_id=vacancy_application.vacancy.id, application_id=vacancy_application.id)