# Editor Document Functions
# Contains create for Document model (for editor type)

import urllib.parse, requests, io, os, logging, platform, shutil, subprocess
from bs4 import BeautifulSoup
from ckeditor_uploader.views import upload as ckeditor_upload
from documents.forms import CreateDocumentForm
from documents.models import Folder, File, Document
from django.http import HttpResponseForbidden
from django.core.exceptions import ValidationError
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, redirect
from django.utils.text import slugify
from docx import Document as DocxDocument
from docx.shared import Inches
from raadaa import settings

logger = logging.getLogger(__name__)

# Formatting content in CKeditor
def add_formatted_content(doc, soup, word_dir):
    """Parse HTML and add formatted content and images to the .docx document."""
    current_paragraph = None

    def add_text_to_paragraph(text, bold=False, italic=False):
        """Helper to add text to the current paragraph with formatting."""
        nonlocal current_paragraph
        if not current_paragraph:
            current_paragraph = doc.add_paragraph()
        run = current_paragraph.add_run(text or '')
        run.bold = bold
        run.italic = italic

    for element in soup.find_all(recursive=False):  # Process top-level elements only
        if element.name == 'h1':
            current_paragraph = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            current_paragraph = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'p':
            current_paragraph = doc.add_paragraph()
            # Process nested elements within <p>
            for child in element.children:
                if child.name == 'strong':
                    add_text_to_paragraph(child.get_text(), bold=True)
                elif child.name == 'em':
                    add_text_to_paragraph(child.get_text(), italic=True)
                elif child.name == 'img':
                    img_src = child.get('src')
                    if img_src:
                        try:
                            parsed_src = urllib.parse.urlparse(img_src)
                            if parsed_src.scheme in ('http', 'https'):
                                print(f"Downloading remote image: {img_src}")
                                response = requests.get(img_src, timeout=5)
                                if response.status_code == 200:
                                    img_data = io.BytesIO(response.content)
                                    doc.add_picture(img_data, width=Inches(4.0))
                                else:
                                    print(f"Failed to download image: {img_src} (Status: {response.status_code})")
                                    current_paragraph = doc.add_paragraph(f"[Image failed to load: {img_src}]")
                            else:
                                # Handle images from CKEditor uploads
                                clean_src = img_src.replace(settings.MEDIA_URL, '').lstrip('/')

                                if settings.DEBUG:
                                    # Local dev: images exist on disk
                                    img_path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, clean_src))
                                    print(f"Attempting to add local image: {img_path}")
                                    if os.path.exists(img_path):
                                        doc.add_picture(img_path, width=Inches(4.0))
                                    else:
                                        print(f"Local image not found: {img_path}")
                                        doc.add_paragraph(f"[Image not found: {img_path}]")
                                else:
                                    # Production: files are on S3 (remote)
                                    file_url = settings.MEDIA_URL + clean_src
                                    print(f"Attempting to fetch image from: {file_url}")
                                    try:
                                        response = requests.get(file_url)
                                        if response.status_code == 200:
                                            image_stream = io.BytesIO(response.content)
                                            doc.add_picture(image_stream, width=Inches(4.0))
                                        else:
                                            print(f"Image not accessible at: {file_url}")
                                            doc.add_paragraph(f"[Image not found: {file_url}]")
                                    except Exception as e:
                                        print(f"Error fetching image: {e}")
                                        doc.add_paragraph(f"[Error loading image: {file_url}]")
                        except Exception as e:
                            print(f"Error adding image {img_src}: {e}")
                            current_paragraph = doc.add_paragraph(f"[Error loading image: {img_src}]")
                else:
                    add_text_to_paragraph(child.get_text() if hasattr(child, 'get_text') else str(child))
        elif element.name == 'ul':
            current_paragraph = None
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='ListBullet')
        elif element.name == 'ol':
            current_paragraph = None
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='ListNumber')
        else:
            current_paragraph = None
            add_text_to_paragraph(element.get_text() if hasattr(element, 'get_text') else str(element))

# Upload media in CKEditor helper function
@login_required
@csrf_exempt  # Required for CKEditor’s POST uploads
def custom_ckeditor_upload(request):
    if hasattr(request, 'tenant') and request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized to perform actions for this company.")
    if not request.user.is_authenticated:
        return HttpResponseForbidden("You must be logged in to upload images.")
    logger.info(f"User {request.user.username} uploading image to CKEditor")
    print(f"User {request.user.username} uploading image to CKEditor")
    response = ckeditor_upload(request)
    if response.status_code == 200:
        logger.info(f"Image upload successful for user {request.user.username}")
        print(f"Image upload successful for user {request.user.username}")
    else:
        logger.error(f"Image upload failed for user {request.user.username}: {response.content}")
        print(f"Image upload failed for user {request.user.username}: {response.content}")
    return response

# Create editor document
@login_required
def create_from_editor(request):
    # Validate that the user belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: User does not belong to this company.")
    
    def documents_word_upload(request):
        tenant_name = request.tenant.name
        username = request.user.username if request.user else "anonymous"
        return os.path.join('documents', tenant_name, username, 'word')

    def documents_pdf_upload(request):
        tenant_name = request.tenant.name
        username = request.user.username if request.user else "anonymous"
        return os.path.join('documents', tenant_name, username, 'pdf')

    if request.method == "POST":
        form = CreateDocumentForm(request.POST)
        if form.is_valid():
            title = form.cleaned_data["title"]
            content = form.cleaned_data["content"]

            # Create a .docx file
            doc = DocxDocument()
            doc.add_heading(title, level=1)

            # Parse HTML content using BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')

            # Define file paths
            if settings.DEBUG:
                word_dir = os.path.join(settings.MEDIA_ROOT, documents_word_upload(request))
                pdf_dir = os.path.join(settings.MEDIA_ROOT, documents_pdf_upload(request))
            else:
                word_dir = f"{settings.MEDIA_URL}{documents_word_upload(request).rstrip('/')}/"
                pdf_dir = f"{settings.MEDIA_URL}{documents_pdf_upload(request).rstrip('/')}/"

            os.makedirs(word_dir, exist_ok=True)
            os.makedirs(pdf_dir, exist_ok=True)

            # Add formatted content and images
            add_formatted_content(doc, soup, word_dir)

            # Get or create the "Template Document" folder
            user = request.user
            department = user.department if hasattr(user, 'department') else None
            team = user.teams.first() if hasattr(user, 'teams') and user.teams.exists() else None

            # Ensure user has a department or team
            if not department and not team:
                messages.error(request, "User must be associated with a department or team.")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Choose department or team for the folder (prefer department if available)
            folder_defaults = {
                'created_by': user,
                'is_public': True,
            }

            try:
                template_folder, created = Folder.objects.get_or_create(
                    tenant=request.tenant,
                    name="Template Document",
                    defaults=folder_defaults
                )
            except ValidationError as e:
                print(f"Folder creation error: {e}")
                messages.error(request, f"Error creating Template Document folder: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Validate folder access
            # if template_folder.department and template_folder.department != user.department:
            #     messages.error(request, "Invalid Department for Template Document folder.")
            #     return render(request, 'documents/create_from_editor.html', {'form': form})
            # if template_folder.team and template_folder.team not in user.teams.all():
            #     messages.error(request, "Invalid Team for Template Document folder.")
            #     return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save the .docx file
            word_filename = f"{slugify(title)}_{request.user.id}_{template_folder.id}.docx"
            word_path = os.path.join(word_dir, word_filename)
            try:
                doc.save(word_path)
                print(f"Saved .docx: {word_path}")
            except Exception as e:
                messages.error(request, f"Error creating .docx file: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Generate and save the .pdf file using LibreOffice
            pdf_filename = f"{slugify(title)}_{request.user.id}_{template_folder.id}.pdf"
            if settings.DEBUG:
                relative_pdf_path = os.path.join(documents_pdf_upload(request), pdf_filename)
            else:
                relative_pdf_path = f"{settings.MEDIA_URL}{documents_pdf_upload(request).rstrip('/')}/{pdf_filename}"
            
            if settings.DEBUG:
                absolute_pdf_path = os.path.join(settings.MEDIA_ROOT, relative_pdf_path)
            else:
                absolute_pdf_path = f"{settings.MEDIA_URL}{relative_pdf_path.rstrip('/')}/"

            # Choose the right LibreOffice path
            if platform.system() == "Windows":
                paths = [
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                    r"C:\Program Files\LibreOffice\program\soffice.exe"
                ]
                libreoffice_path = next((p for p in paths if os.path.exists(p)), None)
            else:
                libreoffice_path = shutil.which("libreoffice") or shutil.which("soffice")

            # Check if LibreOffice exists
            if not libreoffice_path or not os.path.exists(libreoffice_path):
                messages.error(request, "LibreOffice not found. Make sure it's installed and in PATH.")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            try:
                print("Starting PDF conversion with LibreOffice")
                abs_word_path = os.path.abspath(word_path)
                abs_output_dir = os.path.dirname(os.path.abspath(absolute_pdf_path))

                # Run LibreOffice to convert .docx to .pdf
                result = subprocess.run(
                    [libreoffice_path, "--headless", "--convert-to", "pdf", "--outdir", abs_output_dir, abs_word_path],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    check=True,
                    timeout=30
                )

                print("LibreOffice path:", libreoffice_path)
                print("abs_word_path:", abs_word_path)
                print("abs_output_dir:", abs_output_dir)
                print("LibreOffice stdout:", result.stdout.decode())
                print("LibreOffice stderr:", result.stderr.decode())

                # Confirm output PDF file exists
                if not os.path.exists(absolute_pdf_path):
                    if os.path.exists(word_path):
                        os.remove(word_path)
                    messages.error(request, "PDF file was not generated.")
                    return render(request, 'documents/create_from_editor.html', {'form': form})

            except subprocess.CalledProcessError as e:
                print(f"LibreOffice conversion error: {e.stderr.decode()}")
                messages.error(request, f"Error converting to PDF: {e.stderr.decode()}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            except Exception as e:
                print(f"Unexpected error during PDF conversion: {e}")
                messages.error(request, f"Unexpected error converting to PDF: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save to File
            try:
                public_word_file = File(
                    tenant=request.tenant,
                    original_name=word_filename,
                    file=os.path.join(documents_word_upload(request), word_filename),
                    folder=template_folder,
                    uploaded_by=user
                )
                public_word_file.save()

                public_pdf_file = File(
                    tenant=request.tenant,
                    original_name=pdf_filename,
                    file=relative_pdf_path,
                    folder=template_folder,
                    uploaded_by=user
                )
                public_pdf_file.save()
            except Exception as e:
                print(f"Error saving to File: {e}")
                messages.error(request, f"Error saving files to public storage: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save to Document
            document = Document(
                document_type='Uploaded',
                document_source='editor',
                company_name='N/A',
                company_address='N/A',
                contact_person_name='N/A',
                contact_person_email='N/A',
                contact_person_designation='N/A',
                sales_rep='N/A',
                created_by=request.user,
                tenant=request.tenant,
                word_file=os.path.join(documents_pdf_upload(request), word_filename),
                pdf_file=relative_pdf_path
            )
            try:
                document.save()
            except Exception as e:
                messages.error(request, f"Error saving document: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            messages.success(request, 'Document created successfully!')
            return redirect("document_list")
        else:
            print("Form errors:", form.errors)
            messages.error(request, 'Please correct the errors below.')
    else:
        form = CreateDocumentForm()
    return render(request, 'documents/create_from_editor.html', {'form': form})