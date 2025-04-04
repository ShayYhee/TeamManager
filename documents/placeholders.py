from datetime import datetime
from docx import Document as DocxDocument

def replace_placeholders(doc, replacements, document_type):
    """
    Replaces placeholders in a Word document while preserving formatting.
    - Ensures 'Company Address' is bold.
    - Formats the date correctly.
    - Replaces text in paragraphs, tables, headers, and footers.
    """

    # Format date based on document type
    today = datetime.now()
    if document_type == "approval":
        day = today.day
        suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
        formatted_date = f"{day}{suffix} {today.strftime('%B, %Y')}"  # e.g., "25th March, 2025"
    else:
        formatted_date = today.strftime("%m/%d/%Y")  # e.g., "03/25/2025"

    # Add formatted date to replacements
    replacements["{{Date}}"] = formatted_date

    # Function to replace text inside runs while keeping formatting
    def replace_text_in_runs(runs):
        for run in runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)
                    
                    # Apply bold formatting only to Company Address
                    if key == "{{Company Name}}":
                        run.bold = True  

    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_text_in_runs(para.runs)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_runs(para.runs)

    # Replace in headers and footers
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            replace_text_in_runs(header_para.runs)
        for footer_para in section.footer.paragraphs:
            replace_text_in_runs(footer_para.runs)

    return doc